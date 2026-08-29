"""
diagnose.py
===========
Run this once on the Windows laptop, from the DataStage_to_SQL folder:

    python diagnose.py

It checks the environment, finds your DataStage exports, parses a real one, and
tries a full DOCX to PDF conversion. It prints a short summary that fits on one
screen, and writes the long version to diagnose_report.txt.

Nothing is modified. No API keys, endpoints or file contents are printed.

Options:
    python diagnose.py --xml "path\\to\\export.xml"    parse a specific export
    python diagnose.py --template "path\\to\\temp.docx"
"""

from __future__ import annotations

import os
import sys
import tempfile
from pathlib import Path

LINES: list[str] = []
DETAIL: list[str] = []


def out(text: str = "") -> None:
    print(text)
    LINES.append(text)


def detail(text: str) -> None:
    DETAIL.append(text)


def head(n: int, title: str) -> None:
    out("")
    out("=" * 64)
    out(f" {n}  {title}")
    out("=" * 64)


def row(label: str, value: str) -> None:
    if len(label) > 15:
        out(f" {label}")
        out(f" {'':<16}{value}")
    else:
        out(f" {label:<16}{value}")


def short(path) -> str:
    try:
        return str(Path(path).resolve().relative_to(Path.cwd()))
    except Exception:
        return str(path)


# ---------------------------------------------------------------- 1 environment

def section_environment() -> None:
    head(1, "ENVIRONMENT")
    v = sys.version_info
    row("python", f"{v.major}.{v.minor}.{v.micro}   platform: {sys.platform}")
    row("cwd", short(Path.cwd()))
    row("in venv", "yes" if sys.prefix != sys.base_prefix else "NO")

    import importlib
    import importlib.metadata as md

    got = []
    for module, dist in [("docx", "python-docx"), ("lxml", "lxml"),
                         ("docx2pdf", "docx2pdf"), ("win32com.client", "pywin32"),
                         ("crewai", "crewai"), ("langchain_openai", "langchain-openai"),
                         ("flask", "flask")]:
        try:
            importlib.import_module(module)
            try:
                got.append(f"{dist} {md.version(dist)}")
            except Exception:
                got.append(f"{dist} ok")
        except Exception:
            got.append(f"{dist} MISSING")
    for i in range(0, len(got), 2):
        row("packages" if i == 0 else "", "   ".join(f"{g:<28}" for g in got[i:i + 2]).rstrip())


def section_word() -> None:
    head(2, "WORD AND FONTS")
    if sys.platform != "win32":
        row("skipped", f"not Windows ({sys.platform})")
        return

    try:
        import ctypes
        pid = ctypes.windll.kernel32.GetCurrentProcessId()
        sess = ctypes.c_ulong()
        ctypes.windll.kernel32.ProcessIdToSessionId(pid, ctypes.byref(sess))
        row("session", f"{sess.value}" + ("  (session 0 - Word COM will fail)"
                                          if sess.value == 0 else "  interactive"))
    except Exception as exc:
        row("session", f"unknown ({exc})")

    try:
        import pythoncom
        import win32com.client
        pythoncom.CoInitialize()
        word = None
        try:
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            row("word com", f"OK   version {word.Version}")
        finally:
            if word is not None:
                try:
                    word.Quit()
                except Exception:
                    pass
            pythoncom.CoUninitialize()
    except Exception as exc:
        row("word com", f"FAILED  {type(exc).__name__}: {str(exc)[:80]}")
        detail(f"Word COM error: {exc!r}")

    fonts_dir = Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts"
    marks = []
    for name, f in (("Arial", "arial.ttf"), ("Georgia", "georgia.ttf"),
                    ("Consolas", "consola.ttf")):
        marks.append(f"{name} {'ok' if (fonts_dir / f).exists() else 'MISSING'}")
    row("fonts", "   ".join(marks))

    try:
        import subprocess
        rule = "D4F940AB-401B-4EFC-AADC-AD5F3C50688A"
        res = subprocess.run(
            ["powershell", "-NoProfile", "-Command",
             "$p=Get-MpPreference; for($i=0;$i -lt $p.AttackSurfaceReductionRules_Ids.Count;$i++)"
             "{'{0}={1}' -f $p.AttackSurfaceReductionRules_Ids[$i],$p.AttackSurfaceReductionRules_Actions[$i]}"],
            capture_output=True, text=True, timeout=60)
        rules = {k.strip().upper(): v.strip()
                 for k, v in (l.split("=", 1) for l in res.stdout.splitlines() if "=" in l)}
        action = rules.get(rule)
        row("defender asr", {"1": "BLOCKING office child processes",
                             "2": "audit only", "6": "warn only"}.get(action, "not blocking"))
    except Exception:
        row("defender asr", "could not read policy")


# ------------------------------------------------------------------- 3 template

def section_template(template_arg: str | None) -> None:
    head(3, "TEMPLATE")
    candidates = [template_arg] if template_arg else []
    try:
        sys.path.insert(0, str(Path.cwd()))
        from Backend.Analyzer import generate_doc as gd          # type: ignore
        candidates.append(gd.TEMPLATE_DOCX_PATH)
    except Exception:
        pass
    candidates += ["temp.docx", "Backend/Analyzer/temp.docx", "templates/temp.docx"]
    candidates += [str(p) for p in Path.cwd().rglob("temp.docx")][:3]

    path = next((Path(c) for c in candidates if c and Path(c).is_file()), None)
    if path is None:
        row("template", "NOT FOUND")
        row("looked in", ", ".join(short(c) for c in candidates if c)[:200])
        return

    row("template", f"{short(path)}   {path.stat().st_size:,} bytes")
    try:
        with open(f"{path}:Zone.Identifier", "r"):
            row("blocked", "YES - marked as downloaded, Word opens it in Protected View")
    except (OSError, ValueError):
        row("blocked", "no")

    try:
        from docx import Document
        doc = Document(str(path))
        sec = doc.sections[0]
        hdr = sec.header.paragraphs[0].text.strip()
        breaks = sum(p._element.xml.count('w:type="page"') for p in doc.paragraphs)
        row("header", repr(hdr)[:60] if hdr else "(none - report pages get no header)")
        row("paragraphs", f"{len(doc.paragraphs)}   page breaks: {breaks}")
        row("page", f"{sec.page_width.inches:.2f} x {sec.page_height.inches:.2f} in   "
                    f"margins L{sec.left_margin.inches:.2f} R{sec.right_margin.inches:.2f}")
        row("usable width", f"{sec.page_width.inches - sec.left_margin.inches - sec.right_margin.inches:.2f} in")
    except Exception as exc:
        row("open", f"FAILED  {exc}")


# --------------------------------------------------------------- 4 project files

def section_layout(xml_arg: str | None) -> Path | None:
    head(4, "PROJECT LAYOUT")
    root = Path.cwd()
    skip = {".git", "env", "venv", ".venv", "node_modules", "__pycache__", "site-packages"}
    found: list[Path] = []
    for p in root.rglob("*.xml"):
        if any(part in skip for part in p.parts):
            continue
        found.append(p)
        if len(found) >= 400:
            break

    row("xml files", str(len(found)))
    folders: dict[str, int] = {}
    for p in found:
        key = short(p.parent)
        folders[key] = folders.get(key, 0) + 1
    for folder, count in sorted(folders.items(), key=lambda kv: -kv[1])[:8]:
        row("", f"{count:>4}  {folder}")
        detail(f"folder {folder}: {count} xml")

    for name in ("uploads", "data", "projects", "Backend/uploads", "static/uploads"):
        if (root / name).is_dir():
            row("dir exists", name)

    if xml_arg and Path(xml_arg).is_file():
        return Path(xml_arg)
    # Prefer something that looks like a DataStage export.
    for p in found:
        try:
            head_text = p.open("r", encoding="utf-8", errors="ignore").read(4000)
        except Exception:
            continue
        if "DSExport" in head_text or "<Job " in head_text or "DSJob" in head_text:
            return p
    return found[0] if found else None


# ---------------------------------------------------------------- 5 route wiring

def section_route() -> None:
    head(5, "ROUTE WIRING")
    candidates = list(Path.cwd().rglob("analyzer.py"))
    target = next((p for p in candidates if "routes" in p.parts), None)
    if target is None:
        row("analyzer.py", "not found")
        return
    row("file", short(target))
    try:
        text = target.read_text(encoding="utf-8", errors="ignore")
    except Exception as exc:
        row("read", f"failed: {exc}")
        return
    # Report which helpers are referenced, not the source itself.
    markers = ["send_file", "send_from_directory", "make_response", "Response",
               "download_name", "attachment_filename", "as_attachment",
               "mimetype", "BytesIO", "document_generate", ".read(", "open("]
    present = [m for m in markers if m in text]
    row("uses", ", ".join(present) if present else "(none of the usual helpers)")
    for line in text.splitlines():
        if "document_generate(" in line and "import" not in line:
            row("call", line.strip()[:70])
            break


def _stub_missing_modules() -> list[str]:
    """
    Insert placeholder modules for anything generate_doc imports at module level
    but that is not installed here. Only genuinely missing names are stubbed;
    real packages are left alone.
    """
    import importlib
    import types

    stubbed: list[str] = []
    plan = {
        "crewai": ("Agent", "Crew", "Process", "Task"),
        "langchain_openai": ("AzureChatOpenAI",),
        "docx2pdf": ("convert",),
        "pythoncom": (),
        "win32com": (),
        "win32com.client": (),
    }
    for name, attrs in plan.items():
        try:
            importlib.import_module(name)
            continue
        except Exception:
            pass
        module = types.ModuleType(name)
        for attr in attrs:
            setattr(module, attr, object)
        sys.modules[name] = module
        stubbed.append(name)
    if "win32com" in sys.modules and "win32com.client" in sys.modules:
        sys.modules["win32com"].client = sys.modules["win32com.client"]   # type: ignore
    return stubbed


# ------------------------------------------------------------- 6 real parse test

def section_parse(xml_path: Path | None) -> None:
    head(6, "REAL EXPORT PARSE")
    if xml_path is None:
        row("export", "no XML file found - pass one with --xml")
        return
    row("export", f"{short(xml_path)}   {xml_path.stat().st_size:,} bytes")

    try:
        from lxml import etree
        tree = etree.parse(str(xml_path), etree.XMLParser(recover=True, huge_tree=True))
        root = tree.getroot()
    except Exception as exc:
        row("xml", f"UNPARSEABLE  {exc}")
        return

    def tag(el) -> str:
        return el.tag.rsplit("}", 1)[-1] if isinstance(el.tag, str) else ""

    row("root element", f"<{tag(root)}>")
    jobs = [e for e in root.iter() if tag(e) == "Job"]
    row("job elements", str(len(jobs)))
    if jobs:
        j = jobs[0]
        row("job attrs", ", ".join(f"{k}={str(v)[:18]}" for k, v in list(j.attrib.items())[:4]))

    records = [e for e in root.iter() if tag(e) == "Record"]
    row("record count", str(len(records)))

    types: dict[str, int] = {}
    for r in records:
        types[r.get("Type") or "(none)"] = types.get(r.get("Type") or "(none)", 0) + 1
    ordered = sorted(types.items(), key=lambda kv: -kv[1])
    for i in range(0, min(len(ordered), 12), 3):
        row("record types" if i == 0 else "",
            "  ".join(f"{k} {v}" for k, v in ordered[i:i + 3]))
    detail("all record types: " + repr(types))

    # Which property names actually carry the name, the type and the link partner.
    prop_names: dict[str, int] = {}
    for r in records:
        for p in r:
            if tag(p) == "Property" and p.get("Name"):
                prop_names[p.get("Name")] = prop_names.get(p.get("Name"), 0) + 1
    top = sorted(prop_names.items(), key=lambda kv: -kv[1])[:15]
    row("top properties", ", ".join(k for k, _ in top[:8]))
    detail("property histogram: " + repr(sorted(prop_names.items(), key=lambda kv: -kv[1])[:60]))

    for wanted in ("Name", "StageType", "Partner"):
        row(f"has '{wanted}'", f"{prop_names.get(wanted, 0)} records" if wanted in prop_names
            else "NOT PRESENT  <-- parser assumption to revisit")

    ids = [r.get("Identifier") for r in records if r.get("Identifier")][:6]
    row("identifiers", ", ".join(ids))

    # Now the parser itself. Unrelated packages are stubbed if absent, so a
    # missing crewai or pywin32 does not hide the parser result - which is the
    # part of this report that matters most.
    out("")
    try:
        sys.path.insert(0, str(Path.cwd()))
        stubbed = _stub_missing_modules()
        if stubbed:
            row("note", "stubbed for this test: " + ", ".join(stubbed))
        try:
            from Backend.Analyzer import generate_doc as G      # type: ignore
        except Exception:
            import generate_doc as G                            # type: ignore
        G.LLM_ENABLED = False
        parsed = G.DataStageParser(xml_path).parse_all_jobs()
        row("PARSER jobs", str(len(parsed)))
        for job in parsed[:2]:
            g = G.analyze(job)
            row("  job", job.job_name[:40])
            row("  stages", f"{len(g.stages)}   links {len(g.links)}")
            row("  roles", f"src {len(g.sources)}  ref {len(g.references)}  "
                           f"xfm {len(g.transformations)}  tgt {len(g.targets)}  "
                           f"unclassified {len(g.data_objects)}")
            row("  paths", f"{len(g.paths)}  truncated {g.paths_truncated}  cycles {g.has_cycles}")
            row("  warnings", str(len(job.warnings)))
            for w in job.warnings[:2]:
                row("", w[:66])
            for s in g.stages[:4]:
                row("", f"{s.identifier:<8} {s.display_name[:22]:<24} "
                        f"{s.display_type[:16]:<18} {s.role}")
            detail(f"job {job.job_name}: " + repr([
                (s.identifier, s.name, s.stage_type, s.role, len(s.inputs), len(s.outputs))
                for s in g.stages]))
    except Exception as exc:
        import traceback
        row("PARSER", f"FAILED  {type(exc).__name__}: {str(exc)[:60]}")
        detail("parser traceback:\n" + traceback.format_exc())


# --------------------------------------------------------------- 7 conversion

def section_conversion() -> None:
    head(7, "DOCX TO PDF")
    if sys.platform != "win32":
        row("skipped", "not Windows")
        return
    try:
        import pythoncom
        from docx import Document
        from docx2pdf import convert
    except Exception as exc:
        row("skipped", f"{exc}")
        return
    with tempfile.TemporaryDirectory(prefix="diag_") as tmp:
        d = Path(tmp) / "probe.docx"
        pdf = Path(tmp) / "probe.pdf"
        doc = Document()
        doc.add_paragraph("conversion probe")
        doc.save(str(d))
        pythoncom.CoInitialize()
        try:
            convert(str(d), str(pdf))
            row("result", f"OK   {pdf.stat().st_size:,} bytes" if pdf.is_file()
                else "Word ran but produced no PDF")
        except Exception as exc:
            row("result", f"FAILED  {type(exc).__name__}: {str(exc)[:70]}")
            detail(f"conversion error: {exc!r}")
        finally:
            pythoncom.CoUninitialize()


# ------------------------------------------------------------------- 8 azure

def section_azure() -> None:
    head(8, "AZURE (values are never printed)")
    names = ("AZURE_OPENAI_DEPLOYMENT", "AZURE_OPENAI_ENDPOINT", "AZURE_OPENAI_API_KEY",
             "AZURE_OPENAI_API_VERSION", "OPENAI_API_VERSION",
             "HTTPS_PROXY", "REQUESTS_CA_BUNDLE", "SSL_CERT_FILE")
    marks = [f"{n.replace('AZURE_OPENAI_', '').replace('_', '')[:11]}"
             f"={'set' if os.environ.get(n) else '-'}" for n in names]
    for i in range(0, len(marks), 4):
        row("env vars" if i == 0 else "", "  ".join(f"{m:<16}" for m in marks[i:i + 4]).rstrip())


def main() -> int:
    import argparse
    ap = argparse.ArgumentParser()
    ap.add_argument("--xml", default=None)
    ap.add_argument("--template", default=None)
    args = ap.parse_args()

    out("DIAGNOSTIC REPORT")
    xml: Path | None = Path(args.xml) if args.xml and Path(args.xml).is_file() else None

    def layout_step() -> None:
        nonlocal xml
        xml = section_layout(args.xml)

    for label, fn in (
        ("environment", lambda: section_environment()),
        ("word", lambda: section_word()),
        ("template", lambda: section_template(args.template)),
        ("layout", layout_step),
        ("route", lambda: section_route()),
        ("parse", lambda: section_parse(xml)),
        ("conversion", lambda: section_conversion()),
        ("azure", lambda: section_azure()),
    ):
        try:
            fn()
        except Exception as exc:
            import traceback
            out(f" ERROR in {label}: {type(exc).__name__}: {exc}")
            detail(f"{label} traceback:\n" + traceback.format_exc())

    out("")
    out("=" * 64)
    report = Path("diagnose_report.txt")
    report.write_text("\n".join(LINES) + "\n\nDETAIL\n" + "\n".join(DETAIL), encoding="utf-8")
    out(f" Long version written to {report.resolve().name}")
    out(" Photograph the sections above. Send diagnose_report.txt if asked.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
