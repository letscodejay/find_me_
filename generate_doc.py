"""
generate_doc.py
===============
DataStage workflow XML  ->  Word document (.docx)  ->  PDF report.

Pipeline
--------
    .xml  ->  DataStageParser  ->  WorkflowGraph  ->  Narrative (CrewAI)  ->  DocxBuilder  ->  PdfConverter

Batch behaviour
---------------
    one input file  producing one report   ->  a single .pdf
    anything more   (several files, or one file holding several jobs)
                                            ->  a .zip containing every .pdf

Layout of this file
-------------------
    SECTION A   Visual formatting tunables      <- edit these to restyle the report
    SECTION B   DataStage parser
    SECTION C   Graph analysis
    SECTION D   Narrative generation (LLM)
    SECTION E   Word document builder
    SECTION F   PDF conversion
    SECTION G   Orchestration and CLI
"""

from __future__ import annotations

import io
import json
import logging
import os
import re
import shutil
import tempfile
import threading
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Iterable, Sequence

from crewai import Agent, Crew, Process, Task
from docx import Document
from langchain_openai import AzureChatOpenAI
from docx2pdf import convert as docx_to_pdf
import pythoncom
import win32com.client

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

log = logging.getLogger("generate_doc")


# =============================================================================
# SECTION A   VISUAL FORMATTING TUNABLES
# =============================================================================
# Everything that controls how the report *looks* lives in this block. The
# template (see TEMPLATE_DOCX_PATH) is a plain Word document; the builder
# overrides its default styles with the values below so the output matches the
# approved draft rather than Word's defaults.

# --- Template -----------------------------------------------------------------
# Path to temp.docx. Page 1 is the cover and is left untouched; the report is
# appended from page 2 onward.
TEMPLATE_DOCX_PATH = r"temp.docx"

# Whether to insert a page break before the report starts.
#   "auto"  detect whether temp.docx already ends on a page break (recommended)
#   True    always insert one - use when the template is cover text with no break
#   False   never insert one - use when the template already ends on a break
# Getting this wrong is silent: too many breaks leave a blank page between the
# cover and the contents, too few start the contents on the cover itself.
INSERT_PAGE_BREAK_AFTER_TEMPLATE = "auto"

# Delete any empty trailing paragraphs the template leaves behind before writing.
TRIM_TEMPLATE_TRAILING_EMPTY_PARAGRAPHS = True

# If the template ends on several page breaks in a row, keep only one, so the
# contents starts on the page immediately after the cover rather than leaving a
# blank page between them.
COLLAPSE_EXTRA_TEMPLATE_PAGE_BREAKS = True

# --- Page ---------------------------------------------------------------------
# The template's own page size is kept by default. Table widths are measured
# from the real section width at build time, so the layout fits whatever size
# temp.docx uses - A4 or Letter - without editing anything here.
# Both default off so the template's cover page is not shifted - page geometry
# belongs to temp.docx. Table widths adapt to whatever it uses.
FORCE_PAGE_SIZE = False
PAGE_WIDTH_IN = 8.27          # only used when FORCE_PAGE_SIZE is True
PAGE_HEIGHT_IN = 11.69
FORCE_PAGE_MARGINS = False    # True applies the four margins below to every page

PAGE_MARGIN_TOP_IN = 0.85
PAGE_MARGIN_BOTTOM_IN = 0.75
PAGE_MARGIN_LEFT_IN = 0.90
PAGE_MARGIN_RIGHT_IN = 0.90
# Fallback only. The builder measures the real width from the document.
CONTENT_WIDTH_IN = PAGE_WIDTH_IN - PAGE_MARGIN_LEFT_IN - PAGE_MARGIN_RIGHT_IN

# --- Palette (hex, no leading #) ----------------------------------------------
COLOR_CORAL = "C4503C"        # section headings
COLOR_CORAL_DEEP = "A63D2B"   # section number eyebrow, table captions
COLOR_CORAL_LINE = "E8C3BA"   # hairline under a section heading
COLOR_INK = "17242E"          # body text
COLOR_INK_DIM = "5D6E78"      # table header text
COLOR_INK_FAINT = "8A979E"    # running header / footer, "not identified" text
COLOR_RULE = "DCDFDC"         # table borders
COLOR_TABLE_HEAD_BG = "F1F0EC"
COLOR_KEY_CELL_BG = "F7F6F2"  # left column of key/value tables
COLOR_DIAGRAM_BG = "F7F6F2"

# --- Typefaces ----------------------------------------------------------------
# Use fonts installed on the machine that runs the conversion, otherwise Word
# substitutes silently and the PDF will not match the draft.
FONT_SANS = "Arial"            # headings
FONT_SERIF = "Georgia"         # descriptive prose
FONT_MONO = "Consolas"         # identifiers, technical names, diagram

# --- Type scale (points) ------------------------------------------------------
SIZE_SECTION_NUMBER = 7.5      # "SECTION 4" eyebrow
SIZE_SECTION_TITLE = 13.0
SIZE_BODY = 9.5                # descriptive prose
SIZE_TABLE = 8.5
SIZE_TABLE_HEAD = 7.5
SIZE_CAPTION = 7.0
SIZE_DIAGRAM = 8.0
SIZE_RUNNING = 7.5             # header / footer
SIZE_TOC = 9.5

# --- Spacing (points) ---------------------------------------------------------
SPACE_BEFORE_SECTION = 16.0
SPACE_AFTER_SECTION_TITLE = 7.0
SPACE_AFTER_TABLE = 4.0
SPACE_AFTER_CAPTION = 12.0
SPACE_AFTER_PARAGRAPH = 5.0
LINE_SPACING_BODY = 1.28
LINE_SPACING_TABLE = 1.12

# --- Rules and borders --------------------------------------------------------
SECTION_RULE_WIDTH_EIGHTHS = 6      # hairline under section heading (1/8 pt units)
TABLE_BORDER_WIDTH_EIGHTHS = 4
TABLE_HEAD_BORDER_WIDTH_EIGHTHS = 6

# --- Table behaviour ----------------------------------------------------------
# Fixed layout is what stops a long technical name widening its column and
# pushing the table past the page margin. Leave this on.
TABLE_FIXED_LAYOUT = True
TABLE_REPEAT_HEADER_ROW = True      # long tables repeat their header across pages
TABLE_ROWS_KEEP_TOGETHER = False    # True stops rows splitting mid-cell
TABLE_CELL_PAD_TWIPS = 60           # ~3pt

# --- Prose --------------------------------------------------------------------
JUSTIFY_DESCRIPTIONS = True         # descriptive prose is justified
UPPERCASE_SECTION_NUMBER = True
SECTION_NUMBER_SPACING_TWIPS = 30   # letter-spacing on the eyebrow

# --- Running header / footer --------------------------------------------------
# "template"  the header and footer in temp.docx are left exactly as they are and
#             appear on every page. Nothing below overwrites them.
# "generate"  build a running header and footer from the settings below instead.
HEADER_FOOTER_MODE = "template"

# In template mode, add a page number only if the template footer has none - a
# multi-page report with no page numbers is hard to refer to. Set False to leave
# the footer completely untouched.
ADD_PAGE_NUMBER_IF_TEMPLATE_HAS_NONE = True

# Used only when HEADER_FOOTER_MODE == "generate".
SHOW_RUNNING_HEADER = True
SHOW_RUNNING_FOOTER = True
RUNNING_HEADER_LEFT = "Workflow Analysis Report"   # right side shows the job name
FOOTER_PAGE_FORMAT = "Page {page} of {total}"      # rendered as live Word fields

# --- Content limits -----------------------------------------------------------
# The model writes variable-length text and the export carries names of any
# length. These caps keep the layout identical regardless of what arrives.
MAX_DESCRIPTION_CHARS = 340         # per table cell; trimmed at a sentence end
MAX_PATH_CELL_CHARS = 260           # a very long path chain in Section 8
MAX_PROSE_PARAGRAPHS = 3            # Section 3 and the Section 8 summary
MAX_PROSE_CHARS = 1500
LONG_TOKEN_BREAK_AFTER = 16         # give Word a break point inside long identifiers
TRUNCATION_MARK = "…"

# Diagram: Consolas at SIZE_DIAGRAM fits about 100 characters across the page.
# Wider diagrams shrink until they fit, then fall back to the derived layout.
MAX_DIAGRAM_LINES = 26
DIAGRAM_MIN_FONT = 6.0
MONO_CHAR_WIDTH_RATIO = 0.55        # Consolas advance width as a fraction of the em

# --- Missing-value wording ----------------------------------------------------
TEXT_NOT_IDENTIFIED = "Not identified from file"
TEXT_NOT_APPLICABLE = "Not applicable"

# --- Report wording -----------------------------------------------------------
REPORT_TITLE = "Workflow Analysis Report"
CONTENTS_HEADING = "Contents"
CONTENTS_EYEBROW = "Table of"
SECTION_TITLES = {
    1: "Workflow Overview",
    2: "Workflow Stages",
    3: "Workflow Architecture",
    4: "Sources",
    5: "Transformations",
    6: "References",
    7: "Targets",
    8: "Data Path",
    9: "Observations",
}

# --- LLM ----------------------------------------------------------------------
LLM_ENABLED = True                  # False -> tables only, no descriptive prose
LLM_TEMPERATURE = 0.0               # documentation must be reproducible
LLM_TIMEOUT_SECONDS = 180
LLM_MAX_RETRIES = 1
# The project keeps its Azure settings in Modules/Libraries.py, so they are read
# from there first and fall back to environment variables. If that module
# already builds an AzureChatOpenAI client, it is reused as-is rather than
# constructing a second one.
LLM_CONFIG_MODULES = (
    "Modules.Libraries", "modules.Libraries", "Modules.libraries",
    "modules.libraries", "Libraries",
)
LLM_INSTANCE_NAMES = ("llm", "LLM", "azure_llm", "chat_model", "client", "openai_client")
LLM_CONFIG_NAMES = {
    "deployment": ("AZURE_OPENAI_DEPLOYMENT", "AZURE_DEPLOYMENT", "DEPLOYMENT_NAME",
                   "AZURE_OPENAI_DEPLOYMENT_NAME", "deployment_name", "MODEL_NAME",
                   "AZURE_DEPLOYMENT_NAME"),
    "endpoint": ("AZURE_OPENAI_ENDPOINT", "AZURE_ENDPOINT", "OPENAI_API_BASE",
                 "AZURE_OPENAI_API_BASE", "azure_endpoint", "ENDPOINT"),
    "api_key": ("AZURE_OPENAI_API_KEY", "AZURE_API_KEY", "OPENAI_API_KEY",
                "AZURE_OPENAI_KEY", "api_key", "API_KEY"),
    "api_version": ("AZURE_OPENAI_API_VERSION", "OPENAI_API_VERSION", "API_VERSION",
                    "api_version"),
}
AZURE_API_VERSION_DEFAULT = "2024-08-01-preview"

# --- Analysis limits ----------------------------------------------------------
MAX_PATHS = 200                     # stop enumerating beyond this
MAX_PATH_DEPTH = 50
MAX_STAGES_FOR_DIAGRAM = 25         # wider workflows omit Figure 1

# --- Output -------------------------------------------------------------------
# Used when the caller does not pass one - the Flask route often does not.
DEFAULT_OUTPUT_DIR = "reports"

# Where uploaded or selected exports live. document_generate() is called with a
# project name and a file selection rather than full paths, so it looks for each
# file under these roots, in order, before giving up. Add your own if the first
# run reports that a file could not be found.
PROJECT_ROOTS = (
    "INPUTS/analyzer_INPUTS", "INPUTS", "uploads", "data", "projects",
    "Backend/uploads", ".",
)
PROJECT_FILE_SEARCH_DEPTH = 4          # how deep to search a root as a last resort

# --- Conversion ---------------------------------------------------------------
KEEP_INTERMEDIATE_DOCX = False      # True leaves the .docx next to the .pdf
PDF_CONVERSION_TIMEOUT_SECONDS = 300


# =============================================================================
# SECTION B   DATASTAGE PARSER
# =============================================================================
# Reads an IBM InfoSphere DataStage XML export into a typed model. This layer
# does no formatting and no graph reasoning - it only turns XML into objects.

# Record Type values, grouped by what they mean. Kept as data so a new DataStage
# variant is a one-line change rather than a new branch.
RECORD_TYPES_JOB = {"DSJob", "JobDefn", "Job"}
RECORD_TYPES_STAGE = {
    "CustomStage", "TransformerStage", "Stage", "ContainerStage", "CContainer",
    "CTrxStage", "ServerStage", "ParallelStage",
}
RECORD_TYPES_INPUT = {"CustomInput", "TrxInput", "CTrxInput", "StageInput", "Input"}
RECORD_TYPES_OUTPUT = {"CustomOutput", "TrxOutput", "CTrxOutput", "StageOutput", "Output"}
RECORD_TYPES_ANNOTATION = {"Annotation", "CAnnotation"}

# Records that describe the export rather than the workflow. StageType records
# are stage *definitions*, not stages on the canvas, and ContainerView is the
# designer's view of the job. Listing either in Section 2 would be noise.
RECORD_TYPES_IGNORED = {"StageType", "ContainerView", "JobView", "TableDef"}

# Stage types whose secondary inputs are reference data rather than a second
# stream of records. Only lookup-style stages qualify: a Join takes two equal
# data inputs and a Merge takes a master plus updates, so neither secondary
# input is reference data, and classifying them as such would move a real source
# out of Section 4 and into Section 6.
REFERENCE_CONSUMER_HINTS = ("lookup",)

# Property names that may hold a record's display name, in order of preference.
NAME_PROPERTY_CANDIDATES = ("Name", "StageName", "LinkName", "Identifier")
# Property names that may hold a stage's type.
TYPE_PROPERTY_CANDIDATES = ("StageType", "Type", "OLEType", "StageTypeName")
# Property names that may point at the port on the other end of a link.
PARTNER_PROPERTY_CANDIDATES = ("Partner", "PartnerLink", "LinkPartner")

# Identifier convention: V0S3P2 -> stage V0S3, port index 2.
_PORT_ID_RE = re.compile(r"^(?P<stage>.*?S\d+)P(?P<port>\d+)$", re.IGNORECASE)

# A Partner value can name the far end in two ways, both seen in real exports:
#   "V0S2P1"           just the port
#   "V0S2|V0S2P1"      the stage and the port, pipe separated
# The second form states the owning stage outright, so it is preferred over
# deriving the owner from the port identifier.
def _parse_partner(value: str | None) -> tuple[str | None, str | None]:
    """Return (stage identifier or None, port identifier or None)."""
    if not value:
        return None, None
    parts = [p.strip() for p in value.replace("\\", "|").split("|") if p.strip()]
    if not parts:
        return None, None
    if len(parts) >= 2:
        return parts[0], parts[-1]
    return None, parts[0]


class ParseError(Exception):
    """The file could not be read as a DataStage export."""


class NoJobsFound(ParseError):
    pass


@dataclass
class Stage:
    identifier: str
    name: str | None                      # None -> not carried in the export
    stage_type: str | None
    properties: dict[str, str] = field(default_factory=dict)
    collections: dict[str, list[dict[str, str]]] = field(default_factory=dict)
    annotation: str | None = None
    # filled during graph analysis
    role: str = "unclassified"            # source | reference | transformation | target
    inputs: list[str] = field(default_factory=list)     # link names
    outputs: list[str] = field(default_factory=list)

    @property
    def display_name(self) -> str:
        return self.name or TEXT_NOT_IDENTIFIED

    @property
    def display_type(self) -> str:
        return self.stage_type or TEXT_NOT_IDENTIFIED

    def prop(self, *names: str, default: str | None = None) -> str | None:
        """First non-empty value among the given property names."""
        for n in names:
            v = self.properties.get(n)
            if v:
                return v
        return default


@dataclass
class Link:
    name: str | None
    from_stage: str                       # stage identifier
    to_stage: str
    from_port: str                        # port record identifier
    to_port: str
    to_port_index: int | None = None      # 1 is normally the primary input
    is_reference: bool = False
    properties: dict[str, str] = field(default_factory=dict)

    @property
    def display_name(self) -> str:
        return self.name or TEXT_NOT_IDENTIFIED


@dataclass
class ParsedJob:
    job_name: str
    source_file: str
    root_record: dict[str, Any]           # the DSJob record
    job_record: dict[str, Any]            # the <Job> element's own attributes
    all_records: dict[str, dict[str, Any]]
    stages: list[Stage]
    links: list[Link]
    annotations: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


class DataStageParser:
    """
    Turns a DataStage XML export into ParsedJob objects - one per job in the file.

    Written to be forgiving about the things that vary between DataStage
    versions (element nesting, property casing, record type names) and strict
    about the things that must not be guessed (which job, whether a link
    actually resolved).
    """

    def __init__(self, file_path: str | Path):
        self.path = Path(file_path)
        if not self.path.is_file():
            raise ParseError(f"File not found: {self.path}")

    # -- entry point ----------------------------------------------------------

    def parse_all_jobs(self) -> list[ParsedJob]:
        root = self._load_xml()
        job_elements = self._find_job_elements(root)
        if not job_elements:
            raise NoJobsFound(
                f"No <Job> elements found in {self.path.name}. "
                "The file may not be a DataStage XML export."
            )
        return [self._parse_job(el) for el in job_elements]

    # -- XML loading ----------------------------------------------------------

    def _load_xml(self):
        # lxml where available (faster, better recovery), else stdlib.
        try:
            from lxml import etree
            parser = etree.XMLParser(recover=True, huge_tree=True)
            tree = etree.parse(str(self.path), parser)
            return tree.getroot()
        except ImportError:
            import xml.etree.ElementTree as ET
            return ET.parse(str(self.path)).getroot()
        except Exception as exc:                       # malformed beyond recovery
            raise ParseError(f"Could not parse {self.path.name}: {exc}") from exc

    @staticmethod
    def _tag(el) -> str:
        """Local tag name, namespace stripped."""
        t = el.tag
        if not isinstance(t, str):
            return ""
        return t.rsplit("}", 1)[-1]

    def _iter_children(self, el, *names: str):
        wanted = {n.lower() for n in names}
        for child in el:
            if self._tag(child).lower() in wanted:
                yield child

    def _find_job_elements(self, root) -> list:
        """Locate <Job> elements wherever the export puts them."""
        found = []
        for el in root.iter():
            if self._tag(el) == "Job":
                found.append(el)
        if found:
            return found
        # Some exports wrap everything in a single implicit job.
        if self._tag(root) in ("DSExport", "DataStageExport"):
            has_records = any(self._tag(e) == "Record" for e in root.iter())
            if has_records:
                return [root]
        return []

    # -- text extraction ------------------------------------------------------

    @staticmethod
    def _text_of(el) -> str:
        """
        Full text of an element including CDATA and nested runs.
        DataStage wraps long property values in CDATA and sometimes splits them.
        """
        try:
            parts = list(el.itertext())
        except Exception:
            parts = [el.text or ""]
        return "".join(p for p in parts if p).strip()

    def _read_properties(self, record_el) -> dict[str, str]:
        props: dict[str, str] = {}
        for prop in self._iter_children(record_el, "Property"):
            name = prop.get("Name") or prop.get("name")
            if not name:
                continue
            value = self._text_of(prop)
            # Keep the first non-empty value if a name repeats.
            if name not in props or (not props[name] and value):
                props[name] = value
        return props

    def _read_collections(self, record_el) -> dict[str, list[dict[str, str]]]:
        collections: dict[str, list[dict[str, str]]] = {}
        for coll in self._iter_children(record_el, "Collection"):
            name = coll.get("Name") or coll.get("name") or "Collection"
            rows: list[dict[str, str]] = []
            for sub in self._iter_children(coll, "SubRecord", "Record"):
                rows.append(self._read_properties(sub))
            collections[name] = rows
        return collections

    # -- job parsing ----------------------------------------------------------

    def _parse_job(self, job_el) -> ParsedJob:
        job_record = {
            "identifier": job_el.get("Identifier") or job_el.get("identifier") or "",
            "date_modified": job_el.get("DateModified") or "",
            "time_modified": job_el.get("TimeModified") or "",
        }

        all_records = self._index_records(job_el)
        root_record = self._find_root_record(all_records)

        job_name = (
            self._first_prop(root_record, *NAME_PROPERTY_CANDIDATES)
            or job_record["identifier"]
            or self.path.stem
        )

        stages = self._build_stages(all_records)
        links, warnings = self._resolve_links(all_records, stages)
        annotations = self._collect_annotations(all_records)

        self._attach_ports(stages, links)

        return ParsedJob(
            job_name=job_name,
            source_file=self.path.name,
            root_record=root_record,
            job_record=job_record,
            all_records=all_records,
            stages=stages,
            links=links,
            annotations=annotations,
            warnings=warnings,
        )

    def _index_records(self, job_el) -> dict[str, dict[str, Any]]:
        """One pass over every <Record>, keyed by Identifier. No interpretation."""
        records: dict[str, dict[str, Any]] = {}
        anon = 0
        for rec in job_el.iter():
            if self._tag(rec) != "Record":
                continue
            ident = rec.get("Identifier") or rec.get("identifier")
            if not ident:
                anon += 1
                ident = f"__anon_{anon}"
            records[ident] = {
                "identifier": ident,
                "type": rec.get("Type") or rec.get("type") or "",
                "properties": self._read_properties(rec),
                "collections": self._read_collections(rec),
            }
        return records

    @staticmethod
    def _first_prop(record: dict[str, Any] | None, *names: str) -> str | None:
        if not record:
            return None
        props = record.get("properties", {})
        for n in names:
            v = props.get(n)
            if v:
                return v
        return None

    def _find_root_record(self, records: dict[str, dict[str, Any]]) -> dict[str, Any]:
        for rec in records.values():
            if rec["type"] in RECORD_TYPES_JOB:
                return rec
        return {"identifier": "", "type": "", "properties": {}, "collections": {}}

    def _build_stages(self, records: dict[str, dict[str, Any]]) -> list[Stage]:
        stages: list[Stage] = []
        for rec in records.values():
            if rec["type"] in RECORD_TYPES_IGNORED:
                continue
            if rec["type"] not in RECORD_TYPES_STAGE:
                continue
            stages.append(
                Stage(
                    identifier=rec["identifier"],
                    name=self._first_prop(rec, *NAME_PROPERTY_CANDIDATES),
                    stage_type=self._first_prop(rec, *TYPE_PROPERTY_CANDIDATES)
                    or rec["type"],
                    properties=dict(rec["properties"]),
                    collections=dict(rec["collections"]),
                )
            )
        # Stable, human-friendly order: by numeric part of the identifier.
        stages.sort(key=lambda s: self._identifier_sort_key(s.identifier))
        return stages

    @staticmethod
    def _identifier_sort_key(ident: str) -> tuple:
        nums = [int(n) for n in re.findall(r"\d+", ident)]
        return (len(nums) == 0, nums, ident)

    # -- link resolution ------------------------------------------------------
    # An edge in DataStage is a *pair* of port records joined by a Partner
    # property. Walking outputs only means every edge is produced exactly once.

    @classmethod
    def _owner_of(cls, port_identifier: str) -> str | None:
        m = _PORT_ID_RE.match(port_identifier)
        return m.group("stage") if m else None

    @classmethod
    def _port_index(cls, port_identifier: str) -> int | None:
        m = _PORT_ID_RE.match(port_identifier)
        return int(m.group("port")) if m else None

    def _resolve_links(
        self, records: dict[str, dict[str, Any]], stages: list[Stage]
    ) -> tuple[list[Link], list[str]]:
        stage_ids = {s.identifier for s in stages}
        links: list[Link] = []
        warnings: list[str] = []

        for rec in records.values():
            if rec["type"] not in RECORD_TYPES_OUTPUT:
                continue

            out_id = rec["identifier"]
            raw_partner = self._first_prop(rec, *PARTNER_PROPERTY_CANDIDATES)
            partner_stage, partner_port = _parse_partner(raw_partner)

            if not partner_port and not partner_stage:
                warnings.append(
                    f"Output port {out_id} declares no partner link; the edge is omitted."
                )
                continue

            # The far stage: a pipe-separated Partner names it outright.
            to_stage = partner_stage or (self._owner_of(partner_port) if partner_port else None)

            # The near stage: the port on the other end usually points back with
            # the same stage|port form, which states this stage outright. That is
            # more reliable than inferring ownership from the identifier, so it is
            # tried first and the identifier is only a fallback.
            from_stage = None
            back = records.get(partner_port or "", {})
            if back:
                back_stage, _ = _parse_partner(
                    self._first_prop(back, *PARTNER_PROPERTY_CANDIDATES)
                )
                if back_stage:
                    from_stage = back_stage
            if not from_stage:
                from_stage = self._owner_of(out_id)

            if not from_stage or not to_stage:
                warnings.append(
                    f"Could not determine the stages either side of the link "
                    f"{out_id} -> {raw_partner}."
                )
                continue
            if from_stage not in stage_ids or to_stage not in stage_ids:
                missing = from_stage if from_stage not in stage_ids else to_stage
                warnings.append(
                    f"Link {out_id} -> {raw_partner} references {missing}, which is not "
                    "documented as a stage in this export. This usually indicates a "
                    "boundary with a shared container."
                )
                continue

            in_rec = records.get(partner_port or "", {})
            name = (
                self._first_prop(rec, *NAME_PROPERTY_CANDIDATES)
                or self._first_prop(in_rec, *NAME_PROPERTY_CANDIDATES)
            )

            links.append(
                Link(
                    name=name,
                    from_stage=from_stage,
                    to_stage=to_stage,
                    from_port=out_id,
                    to_port=partner_port or "",
                    to_port_index=self._port_index(partner_port) if partner_port else None,
                    properties=dict(rec["properties"]),
                )
            )

        self._flag_reference_links(links, stages)
        return links, warnings

    def _flag_reference_links(self, links: list[Link], stages: list[Stage]) -> None:
        """
        Mark links that feed a lookup-style stage on a non-primary input port.

        DataStage carries no explicit "this is reference data" flag on the link,
        so this is a derivation, not a reading. Two signals are used, in order:
        an explicit LinkType property where the export provides one, otherwise
        the port index on a lookup/join/merge consumer.
        """
        by_id = {s.identifier: s for s in stages}
        for link in links:
            explicit = link.properties.get("LinkType", "").strip().lower()
            if explicit in ("reference", "2"):
                link.is_reference = True
                continue
            target = by_id.get(link.to_stage)
            if not target or not target.stage_type:
                continue
            stype = target.stage_type.lower()
            if not any(h in stype for h in REFERENCE_CONSUMER_HINTS):
                continue
            if link.to_port_index is not None and link.to_port_index > 1:
                link.is_reference = True

    def _attach_ports(self, stages: list[Stage], links: list[Link]) -> None:
        by_id = {s.identifier: s for s in stages}
        for link in links:
            if link.from_stage in by_id:
                by_id[link.from_stage].outputs.append(link.display_name)
            if link.to_stage in by_id:
                by_id[link.to_stage].inputs.append(link.display_name)

    def _collect_annotations(self, records: dict[str, dict[str, Any]]) -> list[str]:
        notes: list[str] = []
        for rec in records.values():
            if rec["type"] not in RECORD_TYPES_ANNOTATION:
                continue
            text = self._first_prop(rec, "AnnotationText", "Text", "Description")
            if text:
                notes.append(text)
        return notes


# =============================================================================
# SECTION C   GRAPH ANALYSIS
# =============================================================================
# Turns stages and links into the derived facts the report is built from:
# roles, path enumeration, and the observations recorded in Section 9.

ROLE_SOURCE = "source"
ROLE_REFERENCE = "reference"
ROLE_TRANSFORMATION = "transformation"
ROLE_TARGET = "target"
ROLE_UNCLASSIFIED = "unclassified"

ROLE_LABELS = {
    ROLE_SOURCE: "Source",
    ROLE_REFERENCE: "Reference",
    ROLE_TRANSFORMATION: "Transformation",
    ROLE_TARGET: "Target",
    ROLE_UNCLASSIFIED: TEXT_NOT_APPLICABLE,
}


@dataclass
class Observation:
    ref: str
    category: str
    obj: str
    text: str


@dataclass
class WorkflowGraph:
    job: ParsedJob
    sources: list[Stage] = field(default_factory=list)
    references: list[Stage] = field(default_factory=list)
    transformations: list[Stage] = field(default_factory=list)
    targets: list[Stage] = field(default_factory=list)
    data_objects: list[Stage] = field(default_factory=list)   # unclassified
    paths: list[list[str]] = field(default_factory=list)      # stage names
    paths_truncated: bool = False
    total_path_count: int = 0
    has_cycles: bool = False
    observations: list[Observation] = field(default_factory=list)

    # -- convenience for the builder -----------------------------------------
    @property
    def stages(self) -> list[Stage]:
        return self.job.stages

    @property
    def links(self) -> list[Link]:
        return self.job.links

    def link_between(self, from_id: str, to_id: str) -> Link | None:
        for l in self.links:
            if l.from_stage == from_id and l.to_stage == to_id:
                return l
        return None

    def outgoing(self, stage: Stage) -> list[Link]:
        return [l for l in self.links if l.from_stage == stage.identifier]

    def incoming(self, stage: Stage) -> list[Link]:
        return [l for l in self.links if l.to_stage == stage.identifier]


def analyze(job: ParsedJob) -> WorkflowGraph:
    graph = WorkflowGraph(job=job)
    _classify_roles(job)
    for s in job.stages:
        if s.role == ROLE_SOURCE:
            graph.sources.append(s)
        elif s.role == ROLE_REFERENCE:
            graph.references.append(s)
        elif s.role == ROLE_TRANSFORMATION:
            graph.transformations.append(s)
        elif s.role == ROLE_TARGET:
            graph.targets.append(s)
        else:
            graph.data_objects.append(s)

    paths, truncated, cycles, total = _enumerate_paths(job, graph)
    graph.paths = paths
    graph.paths_truncated = truncated
    graph.has_cycles = cycles
    graph.total_path_count = total

    graph.observations = _build_observations(job, graph)
    return graph


def _classify_roles(job: ParsedJob) -> None:
    """
    Roles are derived from link degree - DataStage carries no source/target flag.

    Order matters. A stage with no links at all is left unclassified rather than
    being called a source, because an isolated object on the canvas is not an
    input to anything.
    """
    in_deg: dict[str, int] = {s.identifier: 0 for s in job.stages}
    out_deg: dict[str, int] = {s.identifier: 0 for s in job.stages}
    for l in job.links:
        if l.from_stage in out_deg:
            out_deg[l.from_stage] += 1
        if l.to_stage in in_deg:
            in_deg[l.to_stage] += 1

    # A source whose every outgoing link is a reference link is reference data.
    for s in job.stages:
        i, o = in_deg[s.identifier], out_deg[s.identifier]

        if i == 0 and o == 0:
            s.role = ROLE_UNCLASSIFIED
        elif i == 0:
            out_links = [l for l in job.links if l.from_stage == s.identifier]
            if out_links and all(l.is_reference for l in out_links):
                s.role = ROLE_REFERENCE
            else:
                s.role = ROLE_SOURCE
        elif o == 0:
            s.role = ROLE_TARGET
        else:
            s.role = ROLE_TRANSFORMATION


def _enumerate_paths(
    job: ParsedJob, graph: WorkflowGraph
) -> tuple[list[list[str]], bool, bool, int]:
    """
    Every walk from a source or reference to a target.

    Three guards, because real workflows fan out and the count grows fast:
    a cycle guard (DataStage permits cycles, so this is a note, not an error),
    a depth guard, and a volume guard that reports the true count when it trips.
    """
    adjacency: dict[str, list[str]] = {s.identifier: [] for s in job.stages}
    for l in job.links:
        if l.from_stage in adjacency:
            adjacency[l.from_stage].append(l.to_stage)

    names = {s.identifier: s.display_name for s in job.stages}
    target_ids = {s.identifier for s in graph.targets}
    starts = [s.identifier for s in graph.sources] + [s.identifier for s in graph.references]

    paths: list[list[str]] = []
    truncated = False
    cycles = False
    total = 0

    def walk(node: str, trail: list[str], on_path: set[str]) -> None:
        nonlocal truncated, cycles, total
        if len(trail) > MAX_PATH_DEPTH:
            truncated = True
            return
        if node in target_ids or not adjacency.get(node):
            total += 1
            if len(paths) < MAX_PATHS:
                paths.append([names.get(n, n) for n in trail])
            else:
                truncated = True
            return
        for nxt in adjacency.get(node, []):
            if nxt in on_path:
                cycles = True
                continue
            walk(nxt, trail + [nxt], on_path | {nxt})

    for start in starts:
        walk(start, [start], {start})

    return paths, truncated, cycles, total


def _build_observations(job: ParsedJob, graph: WorkflowGraph) -> list[Observation]:
    """Section 9. Records both exceptions and the fact that checks were run."""
    obs: list[Observation] = []

    def add(category: str, obj: str, text: str) -> None:
        obs.append(Observation(f"OBS-{len(obs) + 1:02d}", category, obj, text))

    for w in job.warnings:
        m = re.search(r"\b([A-Za-z0-9_]*S\d+(?:P\d+)?)\b", w)
        add("Unresolved reference", m.group(1) if m else TEXT_NOT_APPLICABLE, w)

    for s in job.stages:
        if s.name is None:
            add(
                "Incomplete metadata",
                s.identifier,
                f"The object does not carry a technical name in this export. It is listed in "
                f"Section 2 as {TEXT_NOT_IDENTIFIED.lower()}, and no role has been assigned to it.",
            )

    for s in graph.data_objects:
        if s.name is not None:
            add(
                "Isolated object",
                s.identifier,
                f"{s.display_name} declares no input or output links and does not participate "
                "in any data path. It is listed in Section 2 but does not appear in Section 3.",
            )

    if graph.has_cycles:
        add(
            "Cyclic route",
            TEXT_NOT_APPLICABLE,
            "One or more cyclic routes were detected during path enumeration. Affected branches "
            "were terminated at the point of repetition, so Section 8 does not represent the "
            "complete set of paths through the workflow.",
        )

    if graph.paths_truncated:
        add(
            "Enumeration limit",
            TEXT_NOT_APPLICABLE,
            f"Path enumeration reached the applicable limit. {graph.total_path_count} paths were "
            f"identified in total and the first {len(graph.paths)} are listed in Section 8.",
        )
    elif not graph.has_cycles:
        add(
            "Completeness",
            TEXT_NOT_APPLICABLE,
            "Path enumeration completed within the applicable limits and no cyclic routes were "
            "detected. Section 8 therefore represents the complete set of paths through the "
            "workflow.",
        )

    return obs


# =============================================================================
# SECTION D   NARRATIVE GENERATION (LLM)
# =============================================================================
# The crew writes explanation *around* parsed facts. It is never the source of a
# value that appears in a table, and the report is generated with tables only if
# the model is unavailable.

@dataclass
class Narrative:
    architecture_diagram: str = ""
    architecture_description: str = ""
    entity_descriptions: dict[str, str] = field(default_factory=dict)   # stage name -> text
    path_explanations: dict[int, str] = field(default_factory=dict)     # 1-based index -> text
    path_summary: str = ""
    unverified_identifiers: list[str] = field(default_factory=list)


def build_factsheet(graph: WorkflowGraph) -> dict[str, Any]:
    """
    A trimmed, token-bounded view of the workflow for the model.

    all_records is deliberately excluded - for a real job it is tens of thousands
    of tokens of port-identifier noise with no explanatory value.
    """
    def stage_entry(s: Stage) -> dict[str, Any]:
        return {
            "name": s.display_name,
            "id": s.identifier,
            "type": s.display_type,
            "role": ROLE_LABELS.get(s.role, s.role),
            "inputs": s.inputs,
            "outputs": s.outputs,
            "key_properties": _interesting_properties(s),
        }

    return {
        "job_name": graph.job.job_name,
        "source_file": graph.job.source_file,
        "counts": {
            "sources": len(graph.sources),
            "references": len(graph.references),
            "transformations": len(graph.transformations),
            "targets": len(graph.targets),
            "total_stages": len(graph.stages) - len(graph.data_objects),
        },
        "sources": [stage_entry(s) for s in graph.sources],
        "references": [stage_entry(s) for s in graph.references],
        "transformations": [stage_entry(s) for s in graph.transformations],
        "targets": [stage_entry(s) for s in graph.targets],
        "links": [
            {
                "name": l.display_name,
                "from": _name_of(graph, l.from_stage),
                "to": _name_of(graph, l.to_stage),
                "reference": l.is_reference,
            }
            for l in graph.links
        ],
        "paths": [{"index": i + 1, "stages": p} for i, p in enumerate(graph.paths)],
        "designer_annotations": graph.job.annotations[:20],
    }


# Properties worth showing a reader, per stage family. Anything not listed stays
# in all_records; dumping every DataStage property produces noise, not detail.
INTERESTING_PROPERTY_NAMES = (
    "Name", "StageType", "TableName", "SelectStatement", "Query", "SQL",
    "FileName", "Filename", "Directory", "ReadMethod", "WriteMode", "WriteMethod",
    "UpdateAction", "InsertAction", "Partitioning", "PartitionType", "SortKey",
    "LookupType", "LookupFail", "Condition", "Constraint", "Derivation",
    "BusinessKey", "SurrogateKey", "CommitInterval", "ArraySize", "DSN",
    "DatabaseName", "SchemaName", "FunnelType",
)


def _interesting_properties(stage: Stage, limit: int = 8) -> dict[str, str]:
    out: dict[str, str] = {}
    for key in INTERESTING_PROPERTY_NAMES:
        val = stage.properties.get(key)
        if val:
            out[key] = val[:300]
        if len(out) >= limit:
            break
    return out


# What a source or target reads from or writes to. Flat properties are checked
# first; modern connector stages (Snowflake, JDBC, ODBC) instead nest their
# configuration inside an XML blob, so that is searched next.
OBJECT_PROPERTY_CANDIDATES = (
    "TableName", "Table_name", "TableNameInput", "Table", "TargetTable",
    "FileName", "Filename", "File", "FilePath", "Path", "Directory",
    "SelectStatement", "Select_Statement", "SQLStatement", "SQL", "Query",
    "Source", "DataSource", "DSN", "DatabaseName", "SchemaName",
)
# Tag names to look for inside a connector's nested XML configuration.
OBJECT_XML_TAGS = (
    "TableName", "Table", "SelectStatement", "SQL", "Query", "FileName",
    "File", "Path", "Schema", "TargetTable",
)
XML_BLOB_PROPERTY_HINTS = ("XMLProperties", "Properties", "StageProperties", "Config")


def stage_object(stage: Stage) -> str | None:
    """The table, file or query a stage reads from or writes to, if recorded."""
    flat = stage.prop(*OBJECT_PROPERTY_CANDIDATES)
    if flat:
        return _tidy_object(flat)

    # Connector stages nest their configuration in an XML string.
    for key, value in stage.properties.items():
        if not value or "<" not in value:
            continue
        if not (any(h.lower() in key.lower() for h in XML_BLOB_PROPERTY_HINTS)
                or value.lstrip().startswith("<")):
            continue
        found = _search_xml_blob(value)
        if found:
            return _tidy_object(found)
    return None


def _search_xml_blob(blob: str) -> str | None:
    try:
        from lxml import etree
        root = etree.fromstring(blob.encode("utf-8", "ignore"),
                                etree.XMLParser(recover=True))
    except Exception:
        return None
    if root is None:
        return None
    wanted = {t.lower() for t in OBJECT_XML_TAGS}
    for el in root.iter():
        tag = el.tag.rsplit("}", 1)[-1] if isinstance(el.tag, str) else ""
        if tag.lower() in wanted:
            text = "".join(el.itertext()).strip()
            if text:
                return text
    return None


def _tidy_object(text: str) -> str:
    """Collapse a multi-line SQL statement so it fits a table cell."""
    text = " ".join(text.split())
    return text if len(text) <= 160 else text[:159] + "\u2026"


def _name_of(graph: WorkflowGraph, stage_id: str) -> str:
    for s in graph.stages:
        if s.identifier == stage_id:
            return s.display_name
    return stage_id


def known_identifiers(graph: WorkflowGraph) -> set[str]:
    known = {s.display_name for s in graph.stages}
    known |= {l.display_name for l in graph.links}
    known |= {s.display_type for s in graph.stages}
    known.add(graph.job.job_name)
    return {k for k in known if k}


def generate_narrative(graph: WorkflowGraph) -> Narrative:
    """
    Runs the crew. Returns an empty Narrative rather than raising - a factual
    report without prose is useful, a failed download is not.
    """
    if not LLM_ENABLED:
        log.info("LLM disabled; generating tables only.")
        return Narrative(architecture_diagram=build_arrow_diagram(graph))

    factsheet = build_factsheet(graph)
    raw: dict[str, Any] = {}
    for attempt in range(LLM_MAX_RETRIES + 1):
        try:
            raw = _run_crew(factsheet)
            break
        except Exception as exc:
            log.warning("Narrative generation attempt %s failed: %s", attempt + 1, exc)
    if not raw:
        log.warning("Narrative generation unavailable; continuing with tables only.")
        return Narrative(architecture_diagram=build_arrow_diagram(graph))

    narrative = _assemble_narrative(raw, graph)
    narrative.unverified_identifiers = _validate_narrative(narrative, graph)
    if narrative.unverified_identifiers:
        log.warning(
            "Narrative mentions identifiers absent from the export: %s",
            ", ".join(narrative.unverified_identifiers),
        )
    return narrative


def _load_config_module():
    """The project's own settings module, if it is importable from here."""
    import importlib

    for name in LLM_CONFIG_MODULES:
        try:
            return importlib.import_module(name)
        except Exception:
            continue
    return None


def _llm_setting(kind: str, module) -> str | None:
    """A setting from the project's config module, else the environment."""
    for attr in LLM_CONFIG_NAMES[kind]:
        if module is not None:
            value = getattr(module, attr, None)
            if isinstance(value, str) and value.strip():
                return value.strip()
        value = os.environ.get(attr)
        if value and value.strip():
            return value.strip()
    return None


def _build_llm():
    module = _load_config_module()

    # If the project already builds a client, reuse it. Constructing a second
    # one risks disagreeing with whatever patches the project applies to it.
    if module is not None:
        for attr in LLM_INSTANCE_NAMES:
            candidate = getattr(module, attr, None)
            if candidate is not None and isinstance(candidate, AzureChatOpenAI):
                log.info("Reusing the AzureChatOpenAI client from %s.%s",
                         module.__name__, attr)
                return candidate

    deployment = _llm_setting("deployment", module)
    if not deployment:
        looked = ", ".join(LLM_CONFIG_NAMES["deployment"][:4])
        raise RuntimeError(
            "No Azure deployment name found. Looked for "
            f"{looked} in {module.__name__ if module else 'the environment'}. "
            "Add the correct attribute name to LLM_CONFIG_NAMES in generate_doc.py."
        )

    kwargs: dict[str, Any] = {
        "azure_deployment": deployment,
        "api_version": _llm_setting("api_version", module) or AZURE_API_VERSION_DEFAULT,
        "temperature": LLM_TEMPERATURE,
        "timeout": LLM_TIMEOUT_SECONDS,
    }
    # Only pass these when found; otherwise langchain reads them from the
    # environment itself, which is the documented behaviour.
    endpoint = _llm_setting("endpoint", module)
    api_key = _llm_setting("api_key", module)
    if endpoint:
        kwargs["azure_endpoint"] = endpoint
    if api_key:
        kwargs["api_key"] = api_key

    log.info(
        "Azure client: deployment=%s endpoint=%s key=%s (from %s)",
        deployment,
        "set" if endpoint else "from environment",
        "set" if api_key else "from environment",
        module.__name__ if module else "environment",
    )
    return AzureChatOpenAI(**kwargs)


GROUNDING_RULES = """
Rules that override every other instruction:
- Use ONLY the names that appear in the factsheet. Never invent a stage, link,
  table, file or column name.
- If a fact is not in the factsheet, write "not specified in the export".
- Write in measured business English. No marketing language, no metaphors, no
  exclamation marks, no bullet fragments - complete sentences only.
- Do not describe what you were asked to do, and do not refer to the factsheet,
  the model or this report in your output.
- Return valid JSON only. No markdown fences, no commentary before or after.
- Plain text inside every JSON string. No markdown: no **bold**, no `backticks`,
  no bullet characters, no headings, no newlines inside a description.
- Respect every length limit given below. Text over the limit is trimmed before
  it reaches the document, so writing more only loses your closing sentence.
"""


def _run_crew(factsheet: dict[str, Any]) -> dict[str, Any]:
    llm = _build_llm()
    facts = json.dumps(factsheet, indent=1, ensure_ascii=False)
    allowed = sorted(
        {e["name"] for k in ("sources", "references", "transformations", "targets")
         for e in factsheet[k]}
        | {l["name"] for l in factsheet["links"]}
    )
    allowed_block = ", ".join(allowed)

    architect = Agent(
        role="ETL Solution Architect",
        goal="Explain how an ETL workflow is put together, accurately and without embellishment.",
        backstory=(
            "You document data integration workflows for enterprise data teams. "
            "Your readers are engineers and analysts who need to understand a job "
            "they did not write."
        ),
        llm=llm, verbose=False, allow_delegation=False,
    )
    documenter = Agent(
        role="Technical Documentation Specialist",
        goal="Describe individual workflow stages in one or two precise sentences each.",
        backstory=(
            "You write reference documentation. You are concise, you never speculate "
            "beyond the metadata you are given, and you keep every entry the same shape."
        ),
        llm=llm, verbose=False, allow_delegation=False,
    )

    architecture_task = Task(
        description=(
            f"{GROUNDING_RULES}\n"
            f"Permitted names: {allowed_block}\n\n"
            f"Workflow factsheet:\n{facts}\n\n"
            "Produce JSON with two keys.\n\n"
            '"diagram": a plain-text arrow diagram of the workflow, laid out left to '
            "right in execution order using only these characters: - | > + and spaces, "
            "plus the stage names. Sources on the left, targets on the right. Show "
            "branches on separate lines. Keep every line under 100 characters. Put the "
            "role in square brackets under each stage name.\n\n"
            f'"description": at most three paragraphs, {MAX_PROSE_CHARS} characters in total, '
            "explaining the flow - where records enter, how they are combined and "
            "transformed, where the flow branches and on what condition, and where "
            "records end up. Separate paragraphs with a blank line."
        ),
        expected_output='JSON object with keys "diagram" and "description".',
        agent=architect,
    )

    # One call for every stage. One Task per stage would re-send the whole
    # context for each one and turn a 40-stage job into 40 round trips.
    entities_task = Task(
        description=(
            f"{GROUNDING_RULES}\n"
            f"Permitted names: {allowed_block}\n\n"
            f"Workflow factsheet:\n{facts}\n\n"
            "Write one description for EVERY stage listed under sources, references, "
            "transformations and targets. One or two sentences each and AT MOST "
            f"{MAX_DESCRIPTION_CHARS} characters, stating what the stage does in the flow "
            "and anything operationally significant in its properties.\n\n"
            'Return a single JSON object mapping each stage name to its description, '
            'for example {"STAGE_NAME": "text", "OTHER_STAGE": "text"}.'
        ),
        expected_output="JSON object mapping every stage name to a description.",
        agent=documenter,
    )

    paths_task = Task(
        description=(
            f"{GROUNDING_RULES}\n"
            f"Permitted names: {allowed_block}\n\n"
            f"Workflow factsheet:\n{facts}\n\n"
            "For each path listed under paths, write one or two sentences and AT MOST "
            f"{MAX_DESCRIPTION_CHARS} characters explaining what that route represents. "
            "Where a path begins at a reference stage, say that it represents data "
            "influence rather than record movement. Then write a closing summary of the "
            f"paths as a whole, at most {MAX_PROSE_CHARS} characters.\n\n"
            'Return JSON: {"paths": {"1": "text", "2": "text"}, "summary": "text"}.'
        ),
        expected_output='JSON object with keys "paths" and "summary".',
        agent=architect,
    )

    crew = Crew(
        agents=[architect, documenter],
        tasks=[architecture_task, entities_task, paths_task],
        process=Process.sequential,
        verbose=False,
    )
    crew.kickoff()

    return {
        "architecture": _json_from_task(architecture_task),
        "entities": _json_from_task(entities_task),
        "paths": _json_from_task(paths_task),
    }


def _task_text(task) -> str:
    """CrewAI has moved this around between versions; try the known shapes."""
    out = getattr(task, "output", None)
    if out is None:
        return ""
    for attr in ("raw", "raw_output", "exported_output", "result"):
        val = getattr(out, attr, None)
        if isinstance(val, str) and val.strip():
            return val
    return str(out)


def _json_from_task(task) -> dict[str, Any]:
    return _loads_loose(_task_text(task))


def _loads_loose(text: str) -> dict[str, Any]:
    """Parse JSON out of a model response that may carry fences or commentary."""
    if not text:
        return {}
    text = text.strip()
    text = re.sub(r"^```(?:json)?\s*|\s*```$", "", text, flags=re.MULTILINE).strip()
    try:
        val = json.loads(text)
        return val if isinstance(val, dict) else {}
    except json.JSONDecodeError:
        pass
    start, end = text.find("{"), text.rfind("}")
    if start != -1 and end > start:
        try:
            val = json.loads(text[start : end + 1])
            return val if isinstance(val, dict) else {}
        except json.JSONDecodeError:
            pass
    return {}


def _assemble_narrative(raw: dict[str, Any], graph: WorkflowGraph) -> Narrative:
    arch = raw.get("architecture", {}) or {}
    ents = raw.get("entities", {}) or {}
    paths = raw.get("paths", {}) or {}

    diagram = str(arch.get("diagram", "") or "").strip("\n")
    if not _diagram_is_usable(diagram, graph):
        log.info("Model diagram unusable; falling back to the derived layout.")
        diagram = build_arrow_diagram(graph)

    path_texts: dict[int, str] = {}
    for key, val in (paths.get("paths") or {}).items():
        try:
            path_texts[int(key)] = str(val).strip()
        except (TypeError, ValueError):
            continue

    return Narrative(
        architecture_diagram=diagram,
        architecture_description=str(arch.get("description", "") or "").strip(),
        entity_descriptions={
            str(k): str(v).strip() for k, v in ents.items() if str(v).strip()
        },
        path_explanations=path_texts,
        path_summary=str(paths.get("summary", "") or "").strip(),
    )


def _diagram_is_usable(diagram: str, graph: WorkflowGraph) -> bool:
    """
    A model-drawn diagram is accepted only if it is actually about this workflow:
    it must mention most of the stages and stay within a printable width.
    """
    if fit_diagram(diagram) is None:                 # width-checked again at build time
        return False
    named = [s.display_name for s in graph.stages if s.name]
    if not named:
        return False
    present = sum(1 for n in named if n in diagram)
    return present >= max(1, int(len(named) * 0.75))


def build_arrow_diagram(graph: WorkflowGraph, width_in: float | None = None) -> str:
    """
    Deterministic fallback layout, used when the model is unavailable or its
    diagram fails validation, so Section 3 always carries a figure.

    One line per link, in execution order. Grouping a stage's targets onto one
    line was tried first and rejected: a stage feeding four others produced a
    170-character line, which no legible font could fit, and the whole figure was
    dropped. A line per link is never wider than two names.

    Long names are shortened to fit rather than losing the figure altogether.
    """
    if not graph.links or len(graph.stages) > MAX_STAGES_FOR_DIAGRAM:
        return ""

    depth: dict[str, int] = {}

    def compute(stage_id: str, seen: frozenset[str]) -> int:
        if stage_id in depth:
            return depth[stage_id]
        incoming = [l.from_stage for l in graph.links if l.to_stage == stage_id]
        incoming = [i for i in incoming if i not in seen]
        d = 0 if not incoming else 1 + max(compute(i, seen | {stage_id}) for i in incoming)
        depth[stage_id] = d
        return d

    for stage in graph.stages:
        compute(stage.identifier, frozenset())

    ordered = sorted(
        graph.links,
        key=lambda l: (depth.get(l.from_stage, 0), _name_of(graph, l.from_stage),
                       _name_of(graph, l.to_stage)),
    )
    rows = [
        (_name_of(graph, l.from_stage), _name_of(graph, l.to_stage), l.is_reference)
        for l in ordered
    ]

    marker = "  (reference)"
    arrow = "  -->  "
    usable = (width_in or CONTENT_WIDTH_IN) - 0.25
    budget = int(72 * usable / (MONO_CHAR_WIDTH_RATIO * DIAGRAM_MIN_FONT))

    name_w = max(len(a) for a, _, _ in rows)
    longest = max(name_w + len(arrow) + len(b) + (len(marker) if r else 0)
                  for _, b, r in rows for a in [""])
    # Shorten names only as far as necessary, and only if necessary.
    allowed = budget - len(arrow) - len(marker)
    if longest > budget and allowed >= 16:
        limit = max(8, allowed // 2)
        rows = [(_ellipsize(a, limit), _ellipsize(b, limit), r) for a, b, r in rows]
        name_w = max(len(a) for a, _, _ in rows)

    return "\n".join(
        f"{a:<{name_w}}{arrow}{b}{marker if r else ''}" for a, b, r in rows
    )


def _ellipsize(text: str, limit: int) -> str:
    return text if len(text) <= limit else text[: limit - 1] + "\u2026"


def _validate_narrative(narrative: Narrative, graph: WorkflowGraph) -> list[str]:
    """
    Identifier-shaped tokens in the prose must exist in the export.

    A model asked to describe an ETL job will produce plausible table and column
    names. In a document that reads as authoritative, that is the failure mode
    worth catching, so anything unrecognised is reported rather than shipped
    silently.
    """
    known_upper = {k.upper() for k in known_identifiers(graph)}
    corpus = " ".join(
        [narrative.architecture_description, narrative.path_summary]
        + list(narrative.entity_descriptions.values())
        + list(narrative.path_explanations.values())
    )
    candidates = set(re.findall(r"\b[A-Z][A-Z0-9]*(?:_[A-Z0-9]+)+\b", corpus))
    return sorted(c for c in candidates if c.upper() not in known_upper)


# =============================================================================
# SECTION E   WORD DOCUMENT BUILDER
# =============================================================================
# temp.docx is a plain Word document. Everything below overrides its defaults so
# the output matches the approved draft rather than Word's built-in styling.

# --- Content normalisation ---------------------------------------------------
# GPT-4o writes markdown, writes longer than it is asked to, and wraps text
# across lines. DataStage carries names of arbitrary length. Everything that
# reaches the page goes through these first, so the layout does not depend on
# either behaving.

# Underscore markdown is deliberately NOT handled. DataStage technical names are
# full of underscores, and treating them as emphasis markers silently rewrites
# EDW_STG.CUSTOMER into EDWSTG.CUSTOMER - corrupting the one thing in the report
# that must be exact. Asterisk markdown only.
_MD_BOLD_RE = re.compile(r"\*\*(.+?)\*\*", re.S)
_MD_ITALIC_RE = re.compile(r"(?<!\*)\*(?!\s|\*)(.+?)(?<!\s)\*(?!\*)", re.S)
_MD_CODE_RE = re.compile(r"`{1,3}([^`]*)`{1,3}", re.S)
_MD_HEADING_RE = re.compile(r"^\s{0,3}#{1,6}\s*", re.M)
_MD_BULLET_RE = re.compile(r"^\s*[-*+]\s+", re.M)
_SENTENCE_END_RE = re.compile(r"[.!?](?=\s|$)")
ZERO_WIDTH_SPACE = "\u200b"


def clean_text(text: Any) -> str:
    """Strip markdown, normalise whitespace, and flatten to a single line.

    A newline inside a run does nothing in Word, so an unflattened multi-line
    description would render with its words run together.
    """
    if text is None:
        return ""
    out = str(text)
    out = _MD_HEADING_RE.sub("", out)
    out = _MD_BULLET_RE.sub("", out)
    out = _MD_CODE_RE.sub(r"\1", out)
    out = _MD_BOLD_RE.sub(r"\1", out)
    out = _MD_ITALIC_RE.sub(r"\1", out)
    out = out.replace("\u00a0", " ")
    # Joining lines needs punctuation, otherwise former bullet points run
    # together into one unreadable sentence.
    lines = [l.strip() for l in out.splitlines() if l.strip()]
    joined = []
    for i, line in enumerate(lines):
        if i < len(lines) - 1 and not line.endswith((".", "!", "?", ":", ";", ",")):
            line += "."
        joined.append(line)
    return re.sub(r"\s+", " ", " ".join(joined)).strip()


def truncate_at_sentence(text: str, limit: int) -> str:
    """Trim to `limit`, preferring the last sentence end so the cell still reads
    as a finished thought rather than a cut-off clause."""
    if not text or len(text) <= limit:
        return text
    window = text[: limit + 1]
    ends = [m.end() for m in _SENTENCE_END_RE.finditer(window)]
    if ends and ends[-1] >= limit * 0.55:
        return window[: ends[-1]].strip()
    cut = window.rsplit(" ", 1)[0].rstrip(" ,;:-")
    return f"{cut}{TRUNCATION_MARK}"


def soften_long_tokens(text: str, after: int = LONG_TOKEN_BREAK_AFTER) -> str:
    """Insert zero-width spaces inside long unbroken tokens.

    Word will not break a word that has no space or hyphen in it. Without a
    break opportunity a long technical name overflows its cell and widens the
    column, which is what pushes a table past the page margin.
    """
    if not text:
        return ""

    def split(match: re.Match) -> str:
        token = match.group(0)
        if len(token) <= after:
            return token
        parts = [token[i : i + after] for i in range(0, len(token), after)]
        return ZERO_WIDTH_SPACE.join(parts)

    return re.sub(r"\S{%d,}" % (after + 1), split, text)


def limit_prose(text: str) -> list[str]:
    """Split narrative into paragraphs and cap both count and total length, so
    a talkative model cannot push the next section onto another page."""
    if not text:
        return []
    blocks = [clean_text(b) for b in re.split(r"\n\s*\n|\n", str(text)) if b.strip()]
    blocks = [b for b in blocks if b][:MAX_PROSE_PARAGRAPHS]
    kept: list[str] = []
    budget = MAX_PROSE_CHARS
    for block in blocks:
        if budget <= 0:
            break
        kept.append(truncate_at_sentence(block, budget) if len(block) > budget else block)
        budget -= len(kept[-1])
    return kept


def fit_diagram(text: str, usable_width_in: float | None = None) -> tuple[str, float] | None:
    """Return the diagram and the font size it fits at, or None if it cannot fit.

    Shrinking is preferred to wrapping: a wrapped monospace diagram is not a
    smaller diagram, it is an unreadable one.
    """
    if not text:
        return None
    lines = [l.rstrip() for l in text.replace("\t", "    ").split("\n")]
    while lines and not lines[-1]:
        lines.pop()
    if not lines or len(lines) > MAX_DIAGRAM_LINES:
        return None

    usable_in = (usable_width_in or CONTENT_WIDTH_IN) - 0.25   # cell padding and borders
    longest = max(len(l) for l in lines)
    size = SIZE_DIAGRAM
    while size >= DIAGRAM_MIN_FONT:
        fits = int(72 * usable_in / (MONO_CHAR_WIDTH_RATIO * size))
        if longest <= fits:
            return "\n".join(lines), size
        size -= 0.5
    return None


def _rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color.upper().lstrip("#"))


def _set_run(
    run,
    font: str | None = None,
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
    caps: bool = False,
    spacing_twips: int | None = None,
) -> None:
    """Set a run's formatting explicitly, including the east-asian and complex
    script slots - setting only run.font.name lets Word substitute silently."""
    if font:
        run.font.name = font
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for slot in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rfonts.set(qn(slot), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color:
        run.font.color.rgb = _rgb(color)
    if caps:
        rpr = run._element.get_or_add_rPr()
        el = OxmlElement("w:caps")
        el.set(qn("w:val"), "1")
        rpr.append(el)
    if spacing_twips:
        rpr = run._element.get_or_add_rPr()
        el = OxmlElement("w:spacing")
        el.set(qn("w:val"), str(spacing_twips))
        rpr.append(el)


def _set_paragraph(
    para,
    space_before: float | None = None,
    space_after: float | None = None,
    line_spacing: float | None = None,
    align=None,
    keep_with_next: bool | None = None,
) -> None:
    pf = para.paragraph_format
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after = Pt(space_after)
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    if align is not None:
        para.alignment = align
    if keep_with_next is not None:
        pf.keep_with_next = keep_with_next


def _paragraph_bottom_border(para, color: str, width_eighths: int) -> None:
    ppr = para._element.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(width_eighths))
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    ppr.append(borders)


def _shade_cell(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _cell_borders(cell, color: str, width_eighths: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in tc_pr.findall(qn("w:tcBorders")):
        tc_pr.remove(existing)
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(width_eighths))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tc_pr.append(borders)


def _cell_margins(cell, twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for edge in ("top", "start", "bottom", "end"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(twips))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tc_pr.append(mar)


def _fixed_table_layout(table) -> None:
    """Pin the column widths. Without this Word auto-fits, and one long
    technical name is enough to widen its column and overflow the page."""
    tbl_pr = table._tbl.tblPr
    for existing in tbl_pr.findall(qn("w:tblLayout")):
        tbl_pr.remove(existing)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)


def _repeat_header_row(row) -> None:
    """Long tables repeat their header across page breaks. Without this the
    second page of a wide stage table is a grid of numbers with no column names."""
    tr_pr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    tr_pr.append(el)


def _cant_split_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:cantSplit")
    tr_pr.append(el)


def _add_field(paragraph, instruction: str) -> None:
    """Insert a live Word field (PAGE, NUMPAGES) - python-docx has no API for these."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._element.append(begin)
    run._element.append(instr)
    run._element.append(end)
    _set_run(run, font=FONT_MONO, size=SIZE_RUNNING, color=COLOR_INK_FAINT)


class DocxBuilder:
    """Writes one WorkflowGraph into a copy of the template."""

    def __init__(self, graph: WorkflowGraph, narrative: Narrative):
        self.graph = graph
        self.narrative = narrative
        self.doc = Document(TEMPLATE_DOCX_PATH)
        self._table_number = 0
        self.content_width = CONTENT_WIDTH_IN      # replaced by the measured width

    # -- template preparation -------------------------------------------------

    def _measure_content_width(self) -> None:
        """Table widths come from the document, not from an assumed page size."""
        section = self.doc.sections[-1]
        try:
            width = (
                section.page_width.inches
                - section.left_margin.inches
                - section.right_margin.inches
            )
        except (TypeError, AttributeError):
            width = CONTENT_WIDTH_IN
        if not width or width < 2:
            log.warning("Template page width looks wrong; falling back to %sin.", CONTENT_WIDTH_IN)
            width = CONTENT_WIDTH_IN
        self.content_width = round(width, 3)

    def _reset_document_defaults(self) -> None:
        """
        Neutralise the template's document defaults.

        Overriding the Normal style is not enough on its own: temp.docx also
        carries w:docDefaults, which any paragraph or table cell that does not
        set a property explicitly inherits from. Resetting both means the report
        looks the same whatever the template was saved with.
        """
        styles_el = self.doc.styles.element
        defaults = styles_el.find(qn("w:docDefaults"))
        if defaults is None:
            defaults = OxmlElement("w:docDefaults")
            styles_el.insert(0, defaults)

        rpr_default = defaults.find(qn("w:rPrDefault"))
        if rpr_default is None:
            rpr_default = OxmlElement("w:rPrDefault")
            defaults.append(rpr_default)
        for old in rpr_default.findall(qn("w:rPr")):
            rpr_default.remove(old)
        rpr = OxmlElement("w:rPr")
        rfonts = OxmlElement("w:rFonts")
        for slot in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rfonts.set(qn(slot), FONT_SERIF)
        rpr.append(rfonts)
        for tag, val in (("w:sz", str(int(SIZE_BODY * 2))),
                         ("w:szCs", str(int(SIZE_BODY * 2))),
                         ("w:color", COLOR_INK)):
            el = OxmlElement(tag)
            el.set(qn("w:val"), val)
            rpr.append(el)
        rpr_default.append(rpr)

        ppr_default = defaults.find(qn("w:pPrDefault"))
        if ppr_default is None:
            ppr_default = OxmlElement("w:pPrDefault")
            defaults.append(ppr_default)
        for old in ppr_default.findall(qn("w:pPr")):
            ppr_default.remove(old)
        ppr = OxmlElement("w:pPr")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), str(int(SPACE_AFTER_PARAGRAPH * 20)))
        ppr.append(spacing)
        ind = OxmlElement("w:ind")
        for slot in ("w:left", "w:right", "w:firstLine"):
            ind.set(qn(slot), "0")
        ppr.append(ind)
        ppr_default.append(ppr)

    def _override_template_defaults(self) -> None:
        """The template ships with Word's default styles. Redefine the ones the
        report uses so its own formatting wins."""
        self._reset_document_defaults()
        styles = self.doc.styles
        normal = styles["Normal"]
        normal.font.name = FONT_SERIF
        normal.font.size = Pt(SIZE_BODY)
        normal.font.color.rgb = _rgb(COLOR_INK)
        rpr = normal.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for slot in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rfonts.set(qn(slot), FONT_SERIF)
        pf = normal.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(SPACE_AFTER_PARAGRAPH)
        pf.line_spacing = LINE_SPACING_BODY

        for section in self.doc.sections:
            if FORCE_PAGE_SIZE:
                section.page_width = Inches(PAGE_WIDTH_IN)
                section.page_height = Inches(PAGE_HEIGHT_IN)
            if FORCE_PAGE_MARGINS:
                section.top_margin = Inches(PAGE_MARGIN_TOP_IN)
                section.bottom_margin = Inches(PAGE_MARGIN_BOTTOM_IN)
                section.left_margin = Inches(PAGE_MARGIN_LEFT_IN)
                section.right_margin = Inches(PAGE_MARGIN_RIGHT_IN)
        self._measure_content_width()

    def _trim_trailing_empty_paragraphs(self) -> None:
        body = self.doc.element.body
        for para in reversed(self.doc.paragraphs):
            if para.text.strip():
                break
            if "w:br" in para._element.xml:      # keep an explicit page break
                break
            body.remove(para._element)

    def _collapse_trailing_page_breaks(self) -> None:
        """
        Keep at most one page break at the end of the template.

        A cover page followed by two breaks is a common way to produce a blank
        page 2 in Word, and it would push the contents to page 3.
        """
        body = self.doc.element.body
        found = []
        for el in reversed(list(body)):
            tag = el.tag.rsplit("}", 1)[-1]
            if tag == "sectPr":
                continue
            if tag != "p":
                break
            if 'w:type="page"' in el.xml and not el.xpath(".//w:t"):
                found.append(el)
                continue
            break
        for extra in found[1:]:                 # found[0] is the last one; keep it
            body.remove(extra)
        if len(found) > 1:
            log.info(
                "Template ended on %s page breaks; kept one so the contents follows "
                "the cover directly.", len(found),
            )

    def _template_ends_with_page_break(self) -> bool:
        """True when the template already breaks to a new page after the cover."""
        body = self.doc.element.body
        for el in reversed(list(body)):
            tag = el.tag.rsplit("}", 1)[-1]
            if tag == "sectPr":
                continue                       # section properties, not content
            if tag == "tbl":
                return False                   # a table is real content
            if tag == "p":
                xml = el.xml
                if 'w:type="page"' in xml or "w:type='page'" in xml:
                    return True                # explicit page break
                if "<w:br" in xml and "sectPr" in xml:
                    return True                # section break acting as one
                if el.xpath(".//w:t"):
                    return False               # text on the last page, no break
                continue                       # empty paragraph, keep looking back
        return False

    def _start_report_page(self) -> None:
        setting = INSERT_PAGE_BREAK_AFTER_TEMPLATE
        if isinstance(setting, str) and setting.lower() == "auto":
            insert = not self._template_ends_with_page_break()
            log.info(
                "Template already ends on a page break; starting the report there."
                if not insert else
                "Template has no page break after the cover; inserting one."
            )
        else:
            insert = bool(setting)

        if insert:
            para = self.doc.add_paragraph()
            para.add_run().add_break(WD_BREAK.PAGE)
            _set_paragraph(para, space_before=0, space_after=0)

    def _build_running_header_footer(self) -> None:
        """
        In template mode the header and footer defined in temp.docx are left
        exactly as they are, so whatever is set there appears on every page.

        Word stores a header once per section, not per page, so a header placed
        in the template shows on the whole document by design - there is nothing
        to copy forward.
        """
        section = self.doc.sections[-1]

        if HEADER_FOOTER_MODE == "template":
            if ADD_PAGE_NUMBER_IF_TEMPLATE_HAS_NONE:
                self._add_page_number_if_absent(section)
            return

        section.different_first_page_header_footer = True   # cover stays clean

        if SHOW_RUNNING_HEADER:
            para = section.header.paragraphs[0]
            para.text = ""
            _set_paragraph(para, space_after=2)
            left = para.add_run(RUNNING_HEADER_LEFT)
            _set_run(left, FONT_MONO, SIZE_RUNNING, color=COLOR_INK_FAINT, caps=True)
            para.add_run("\t\t")
            right = para.add_run(self.graph.job.job_name)
            _set_run(right, FONT_MONO, SIZE_RUNNING, color=COLOR_INK_FAINT, caps=True)
            _paragraph_bottom_border(para, COLOR_RULE, 4)

        if SHOW_RUNNING_FOOTER:
            para = section.footer.paragraphs[0]
            para.text = ""
            _set_paragraph(para, space_before=2)
            src = para.add_run(self.graph.job.source_file)
            _set_run(src, FONT_MONO, SIZE_RUNNING, color=COLOR_INK_FAINT)
            para.add_run("\t\t")
            self._append_page_fields(para)

    @staticmethod
    def _append_page_fields(para) -> None:
        before, _, after = FOOTER_PAGE_FORMAT.partition("{page}")
        mid, _, tail = after.partition("{total}")
        for text, field in ((before, "PAGE"), (mid, "NUMPAGES")):
            if text:
                run = para.add_run(text)
                _set_run(run, FONT_MONO, SIZE_RUNNING, color=COLOR_INK_FAINT)
            _add_field(para, field)
        if tail:
            run = para.add_run(tail)
            _set_run(run, FONT_MONO, SIZE_RUNNING, color=COLOR_INK_FAINT)

    def _add_page_number_if_absent(self, section) -> None:
        """Only touches the template footer when it carries no page number."""
        footer = section.footer
        if "PAGE" in footer._element.xml.upper():
            return
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        if para.text.strip():
            para.add_run("\t\t")
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._append_page_fields(para)

    # -- primitives -----------------------------------------------------------

    def _heading(self, eyebrow: str, title: str, first: bool = False) -> None:
        eb = self.doc.add_paragraph()
        _set_paragraph(
            eb,
            space_before=0 if first else SPACE_BEFORE_SECTION,
            space_after=1,
            line_spacing=1.0,
            keep_with_next=True,
        )
        run = eb.add_run(eyebrow.upper() if UPPERCASE_SECTION_NUMBER else eyebrow)
        _set_run(
            run, FONT_MONO, SIZE_SECTION_NUMBER, bold=True,
            color=COLOR_CORAL_DEEP, spacing_twips=SECTION_NUMBER_SPACING_TWIPS,
        )

        tp = self.doc.add_paragraph()
        _set_paragraph(
            tp, space_before=0, space_after=SPACE_AFTER_SECTION_TITLE,
            line_spacing=1.0, keep_with_next=True,
        )
        trun = tp.add_run(title)
        _set_run(trun, FONT_SANS, SIZE_SECTION_TITLE, bold=True, color=COLOR_CORAL)
        _paragraph_bottom_border(tp, COLOR_CORAL_LINE, SECTION_RULE_WIDTH_EIGHTHS)

    def _prose(self, text: str) -> None:
        blocks = limit_prose(text)
        if not blocks:
            return
        for block in blocks:
            para = self.doc.add_paragraph()
            _set_paragraph(
                para,
                space_after=SPACE_AFTER_PARAGRAPH,
                line_spacing=LINE_SPACING_BODY,
                align=WD_ALIGN_PARAGRAPH.JUSTIFY if JUSTIFY_DESCRIPTIONS else None,
            )
            run = para.add_run(block)
            _set_run(run, FONT_SERIF, SIZE_BODY, color=COLOR_INK)

    def _caption(self, text: str) -> None:
        self._table_number += 1
        para = self.doc.add_paragraph()
        _set_paragraph(para, space_before=SPACE_AFTER_TABLE, space_after=SPACE_AFTER_CAPTION)
        run = para.add_run(f"Table {self._table_number} · {text}")
        _set_run(run, FONT_MONO, SIZE_CAPTION, bold=True, color=COLOR_CORAL_DEEP, caps=True)

    def _figure_caption(self, text: str) -> None:
        para = self.doc.add_paragraph()
        _set_paragraph(para, space_before=SPACE_AFTER_TABLE, space_after=SPACE_AFTER_CAPTION)
        run = para.add_run(text)
        _set_run(run, FONT_MONO, SIZE_CAPTION, bold=True, color=COLOR_CORAL_DEEP, caps=True)

    def _monospace_block(self, text: str, size: float = SIZE_DIAGRAM) -> None:
        """The arrow diagram. Rendered as text so it stays selectable in the PDF
        and needs no image toolchain - python-docx cannot place SVG."""
        table = self.doc.add_table(rows=1, cols=1)
        table.autofit = False
        if TABLE_FIXED_LAYOUT:
            _fixed_table_layout(table)
        cell = table.cell(0, 0)
        _shade_cell(cell, COLOR_DIAGRAM_BG)
        _cell_borders(cell, COLOR_RULE, TABLE_BORDER_WIDTH_EIGHTHS)
        _cell_margins(cell, 140)
        cell.width = Inches(self.content_width)
        para = cell.paragraphs[0]
        _set_paragraph(para, space_before=0, space_after=0, line_spacing=1.15)
        for i, line in enumerate(text.split("\n")):
            if i:
                para.add_run().add_break()
            run = para.add_run(line.replace("\t", "    "))
            _set_run(run, FONT_MONO, size, color=COLOR_INK)

    def _table(
        self,
        headers: Sequence[str] | None,
        rows: Sequence[Sequence[Any]],
        widths: Sequence[float],
        mono_columns: Iterable[int] = (),
        desc_columns: Iterable[int] = (),
        key_column: int | None = None,
        limits: dict[int, int] | None = None,
    ):
        """
        rows may contain plain strings or (text, style) tuples where style is
        'na' for a missing-value cell.

        `limits` caps a column's characters. Description columns default to
        MAX_DESCRIPTION_CHARS so a long model response cannot change the layout.
        """
        mono_columns = set(mono_columns)
        desc_columns = set(desc_columns)
        limits = dict(limits or {})
        for col in desc_columns:
            limits.setdefault(col, MAX_DESCRIPTION_CHARS)
        n_cols = len(widths)
        table = self.doc.add_table(rows=0, cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        if TABLE_FIXED_LAYOUT:
            _fixed_table_layout(table)

        def fill(cells, values, header: bool) -> None:
            for idx, (cell, value) in enumerate(zip(cells, values)):
                cell.width = Inches(widths[idx])
                _cell_borders(
                    cell, COLOR_RULE,
                    TABLE_HEAD_BORDER_WIDTH_EIGHTHS if header else TABLE_BORDER_WIDTH_EIGHTHS,
                )
                _cell_margins(cell, TABLE_CELL_PAD_TWIPS)

                text, kind = (value if isinstance(value, tuple) else (value, None))
                text = clean_text(text)
                if not header and idx in limits:
                    text = truncate_at_sentence(text, limits[idx])
                text = soften_long_tokens(text)

                para = cell.paragraphs[0]
                _set_paragraph(
                    para, space_before=0, space_after=0,
                    line_spacing=LINE_SPACING_TABLE,
                    align=(
                        WD_ALIGN_PARAGRAPH.JUSTIFY
                        if (JUSTIFY_DESCRIPTIONS and idx in desc_columns and not header)
                        else None
                    ),
                )
                run = para.add_run(text)

                if header:
                    _shade_cell(cell, COLOR_TABLE_HEAD_BG)
                    _set_run(run, FONT_MONO, SIZE_TABLE_HEAD, bold=True,
                             color=COLOR_INK_DIM, caps=True)
                elif kind == "na":
                    _set_run(run, FONT_SERIF, SIZE_TABLE, italic=True, color=COLOR_INK_FAINT)
                elif key_column is not None and idx == key_column:
                    _shade_cell(cell, COLOR_KEY_CELL_BG)
                    _set_run(run, FONT_MONO, SIZE_TABLE_HEAD, color=COLOR_INK_DIM, caps=True)
                elif idx in mono_columns:
                    _set_run(run, FONT_MONO, SIZE_TABLE, color=COLOR_INK)
                else:
                    _set_run(run, FONT_SERIF, SIZE_TABLE, color=COLOR_INK)

        if headers:
            row = table.add_row()
            fill(row.cells, headers, header=True)
            if TABLE_REPEAT_HEADER_ROW:
                _repeat_header_row(row)

        for values in rows:
            row = table.add_row()
            if TABLE_ROWS_KEEP_TOGETHER:
                _cant_split_row(row)
            fill(row.cells, values, header=False)

        return table

    # -- sections -------------------------------------------------------------

    def build(self) -> Document:
        self._override_template_defaults()
        if TRIM_TEMPLATE_TRAILING_EMPTY_PARAGRAPHS:
            self._trim_trailing_empty_paragraphs()
        if COLLAPSE_EXTRA_TEMPLATE_PAGE_BREAKS:
            self._collapse_trailing_page_breaks()
        self._start_report_page()
        self._build_running_header_footer()

        self._section_contents()
        self._section_1_overview()
        self._section_2_stages()
        self._section_3_architecture()
        self._section_4_sources()
        self._section_5_transformations()
        self._section_6_references()
        self._section_7_targets()
        self._section_8_paths()
        self._section_9_observations()
        return self.doc

    def _section_contents(self) -> None:
        self._heading(CONTENTS_EYEBROW, CONTENTS_HEADING, first=True)
        for number, title in SECTION_TITLES.items():
            para = self.doc.add_paragraph()
            _set_paragraph(para, space_before=0, space_after=4, line_spacing=1.0)
            label = para.add_run(f"Section {number}")
            _set_run(label, FONT_MONO, SIZE_CAPTION, bold=True,
                     color=COLOR_CORAL_DEEP, caps=True)
            para.add_run("   ")
            name = para.add_run(title)
            _set_run(name, FONT_SERIF, SIZE_TOC, color=COLOR_INK)

    def _section_1_overview(self) -> None:
        g = self.graph
        self._heading("Section 1", SECTION_TITLES[1])
        rows = [
            ("Filename", g.job.source_file),
            ("# of sources", str(len(g.sources))),
            ("# of targets", str(len(g.targets))),
            ("# of references", str(len(g.references))),
            ("# of transformations", str(len(g.transformations))),
            (
                "# of total stages",
                str(len(g.sources) + len(g.references) + len(g.transformations) + len(g.targets)),
            ),
        ]
        self._table(None, rows, [2.20, self.content_width - 2.20],
                    mono_columns=[1], key_column=0)
        self._caption("Workflow overview")

    def _section_2_stages(self) -> None:
        self._heading("Section 2", SECTION_TITLES[2])
        rows = []
        for s in self.graph.stages:
            role = ROLE_LABELS.get(s.role, s.role)
            rows.append([
                s.identifier,
                s.display_name if s.name else (TEXT_NOT_IDENTIFIED, "na"),
                s.display_type,
                str(len(s.inputs)),
                str(len(s.outputs)),
                role if s.role != ROLE_UNCLASSIFIED else (TEXT_NOT_APPLICABLE, "na"),
            ])
        self._table(
            ["ID", "Technical name", "Type", "# input links", "# output links", "Role"],
            rows,
            [0.70, 1.65, 1.40, 0.62, 0.62, self.content_width - 4.99],
            mono_columns=[0, 1, 2, 3, 4],
        )
        self._caption("Workflow stages and data objects")

    def _section_3_architecture(self) -> None:
        self._heading("Section 3", SECTION_TITLES[3])
        fitted = fit_diagram(self.narrative.architecture_diagram, self.content_width)
        if fitted is None:
            fitted = fit_diagram(build_arrow_diagram(self.graph, self.content_width),
                                 self.content_width)
        if fitted:
            diagram, size = fitted
            self._monospace_block(diagram, size)
            self._figure_caption("Figure 1 · Workflow architecture")
        else:
            # Reached when the workflow is too wide to draw, or when the stage
            # names are long enough that no legible font fits the page.
            self._prose(
                f"The workflow contains {len(self.graph.stages)} stages and cannot be "
                "presented legibly as a diagram at the width of this page. Section 2 lists "
                "every stage and its links, and Section 8 lists the paths between them."
            )
        self._prose(self.narrative.architecture_description)

    def _entity_description(self, stage: Stage) -> Any:
        text = self.narrative.entity_descriptions.get(stage.display_name, "").strip()
        return text if text else (TEXT_NOT_IDENTIFIED, "na")

    def _section_4_sources(self) -> None:
        self._heading("Section 4", SECTION_TITLES[4])
        rows = []
        for s in self.graph.sources:
            obj = stage_object(s)
            out_links = ", ".join(s.outputs) or TEXT_NOT_APPLICABLE
            rows.append([
                s.display_name, s.identifier, s.display_type,
                obj if obj else (TEXT_NOT_IDENTIFIED, "na"),
                out_links,
                self._entity_description(s),
            ])
        if not rows:
            rows = [[(TEXT_NOT_IDENTIFIED, "na"), "", "", "", "", ""]]
        self._table(
            ["Technical name", "ID", "Type", "Object", "Output link", "Description"],
            rows,
            [1.15, 0.50, 1.00, 1.25, 0.90, self.content_width - 4.80],
            mono_columns=[0, 1, 2, 3, 4], desc_columns=[5],
        )
        self._caption("Sources")

    def _section_5_transformations(self) -> None:
        self._heading("Section 5", SECTION_TITLES[5])
        rows = []
        for s in self.graph.transformations:
            rows.append([
                s.display_name, s.identifier, s.display_type,
                f"{len(s.inputs)} / {len(s.outputs)}",
                self._entity_description(s),
            ])
        if not rows:
            rows = [[(TEXT_NOT_IDENTIFIED, "na"), "", "", "", ""]]
        self._table(
            ["Technical name", "ID", "Type", "In / Out", "Description"],
            rows,
            [1.15, 0.50, 1.05, 0.60, self.content_width - 3.30],
            mono_columns=[0, 1, 2, 3], desc_columns=[4],
        )
        self._caption("Transformations")

    def _section_6_references(self) -> None:
        self._heading("Section 6", SECTION_TITLES[6])
        rows = []
        for s in self.graph.references:
            used_by = ", ".join(
                sorted({_name_of(self.graph, l.to_stage) for l in self.graph.outgoing(s)})
            )
            rows.append([
                s.display_name, s.identifier, s.display_type,
                used_by if used_by else (TEXT_NOT_IDENTIFIED, "na"),
                self._entity_description(s),
            ])
        if not rows:
            rows = [[(TEXT_NOT_APPLICABLE, "na"), "", "", "", ""]]
        self._table(
            ["Technical name", "ID", "Type", "Used by", "Description"],
            rows,
            [1.15, 0.50, 1.10, 1.00, self.content_width - 3.75],
            mono_columns=[0, 1, 2, 3], desc_columns=[4],
        )
        self._caption("References")

    def _section_7_targets(self) -> None:
        self._heading("Section 7", SECTION_TITLES[7])
        rows = []
        for s in self.graph.targets:
            obj = stage_object(s)
            mode = s.prop("WriteMode", "WriteMethod", "UpdateAction")
            rows.append([
                s.display_name, s.identifier, s.display_type,
                obj if obj else (TEXT_NOT_IDENTIFIED, "na"),
                mode if mode else (TEXT_NOT_IDENTIFIED, "na"),
                self._entity_description(s),
            ])
        if not rows:
            rows = [[(TEXT_NOT_IDENTIFIED, "na"), "", "", "", "", ""]]
        self._table(
            ["Technical name", "ID", "Type", "Object", "Write mode", "Description"],
            rows,
            [1.15, 0.50, 1.00, 1.20, 0.75, self.content_width - 4.60],
            mono_columns=[0, 1, 2, 3], desc_columns=[5],
        )
        self._caption("Targets")

    def _section_8_paths(self) -> None:
        self._heading("Section 8", SECTION_TITLES[8])
        rows = []
        for i, path in enumerate(self.graph.paths, start=1):
            explanation = self.narrative.path_explanations.get(i, "").strip()
            rows.append([
                str(i),
                " → ".join(path),
                str(len(path)),
                explanation if explanation else (TEXT_NOT_IDENTIFIED, "na"),
            ])
        if not rows:
            rows = [[(TEXT_NOT_APPLICABLE, "na"), "", "", ""]]
        self._table(
            ["#", "Path", "Stages in path", "Explanation"],
            rows,
            [0.32, 2.35, 0.50, self.content_width - 3.17],
            mono_columns=[0, 1, 2], desc_columns=[3],
            limits={1: MAX_PATH_CELL_CHARS},
        )
        self._caption("Data paths")
        self._prose(self.narrative.path_summary)

    def _section_9_observations(self) -> None:
        self._heading("Section 9", SECTION_TITLES[9])
        rows = []
        for o in self.graph.observations:
            rows.append([
                o.ref, o.category,
                o.obj if o.obj != TEXT_NOT_APPLICABLE else (TEXT_NOT_APPLICABLE, "na"),
                o.text,
            ])
        for ident in self.narrative.unverified_identifiers:
            rows.append([
                f"OBS-{len(rows) + 1:02d}", "Unverified reference", ident,
                f"The descriptive text refers to {ident}, which does not appear in the export. "
                "The reference could not be verified against the source file.",
            ])
        if not rows:
            rows = [["OBS-01", "Completeness", (TEXT_NOT_APPLICABLE, "na"),
                     "No exceptions were recorded during analysis."]]
        self._table(
            ["Ref", "Category", "Object", "Observation"],
            rows,
            [0.68, 1.10, 0.85, self.content_width - 2.63],
            mono_columns=[0, 2], desc_columns=[3],
        )
        self._caption("Observations")


# =============================================================================
# SECTION F   PDF CONVERSION
# =============================================================================
# docx2pdf drives Microsoft Word through COM. Three things break in production,
# in this order: COM is not initialised on the worker thread, Word is not
# concurrent, and a crash leaves WINWORD.EXE resident.

_WORD_LOCK = threading.Lock()


def convert_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    """
    Convert with Word via COM. Isolated behind one function so that swapping in
    LibreOffice for a non-Windows deployment is a single-file change:

        soffice --headless --convert-to pdf --outdir <dir> <docx>
    """
    # COM must be initialised per thread. Flask serves each request on a worker
    # thread, so the main thread's initialisation does not apply - this is the
    # "CoInitialize has not been called" failure.
    pythoncom.CoInitialize()
    try:
        # Two simultaneous conversions on one Word instance hang or corrupt the
        # output, so conversion is serialised.
        with _WORD_LOCK:
            docx_to_pdf(str(docx_path), str(pdf_path))
    finally:
        _quit_stray_word()
        pythoncom.CoUninitialize()

    if not pdf_path.is_file():
        raise RuntimeError(f"Word did not produce {pdf_path.name}.")


def _quit_stray_word() -> None:
    """A crash mid-conversion leaves Word resident; enough of them exhaust the box."""
    try:
        word = win32com.client.GetActiveObject("Word.Application")
        if word.Documents.Count == 0:
            word.Quit()
    except Exception:
        pass          # no running instance, or it is in use - nothing to clean up


# =============================================================================
# SECTION G   ORCHESTRATION AND CLI
# =============================================================================

_SAFE_NAME_RE = re.compile(r"[^A-Za-z0-9._-]+")


def _safe_name(text: str) -> str:
    return _SAFE_NAME_RE.sub("_", text).strip("_") or "report"


def _unique_stem(stem: str, used: set[str]) -> str:
    """
    Two uploads can carry the same filename, and two files in different folders
    can share a stem. Without this the second report overwrites the first on
    disk and shadows it inside the ZIP.
    """
    candidate = stem
    n = 2
    while candidate.lower() in used:
        candidate = f"{stem}_{n}"
        n += 1
    used.add(candidate.lower())
    return candidate


def build_report(job: ParsedJob, out_dir: Path, stem: str) -> Path:
    """Analyse one job, write the .docx, convert it, and return the .pdf path."""
    graph = analyze(job)
    narrative = generate_narrative(graph)

    docx_path = out_dir / f"{stem}.docx"
    pdf_path = out_dir / f"{stem}.pdf"

    DocxBuilder(graph, narrative).build().save(str(docx_path))
    log.info("Wrote %s", docx_path.name)

    # Convert in a temp directory and move the result. Word's file locking makes
    # conversion directly onto a network path unreliable.
    with tempfile.TemporaryDirectory(prefix="wf_report_") as tmp:
        tmp_dir = Path(tmp)
        tmp_docx = tmp_dir / docx_path.name
        tmp_pdf = tmp_dir / pdf_path.name
        shutil.copy2(docx_path, tmp_docx)
        convert_to_pdf(tmp_docx, tmp_pdf)
        shutil.move(str(tmp_pdf), str(pdf_path))

    if not KEEP_INTERMEDIATE_DOCX:
        docx_path.unlink(missing_ok=True)

    log.info("Wrote %s", pdf_path.name)
    return pdf_path


def build_reports_for_file(
    xml_path: Path, out_dir: Path, used_names: set[str] | None = None
) -> list[Path]:
    """
    One PDF per job. A file holding several jobs produces several reports rather
    than silently documenting the first one.
    """
    used_names = used_names if used_names is not None else set()
    jobs = DataStageParser(xml_path).parse_all_jobs()
    stem = _safe_name(xml_path.stem)
    pdfs: list[Path] = []

    for job in jobs:
        if not job.stages:
            log.warning(
                "%s / %s contains no stages; no report generated.",
                xml_path.name, job.job_name,
            )
            continue
        base = stem if len(jobs) == 1 else f"{stem}__{_safe_name(job.job_name)}"
        pdfs.append(build_report(job, out_dir, _unique_stem(base, used_names)))

    return pdfs


def generate(inputs: Sequence[str | Path], out_dir: str | Path) -> Path:
    """
    Main entry point, including for a Flask route.

    Returns a single .pdf when exactly one report was produced, otherwise a .zip
    containing every report.
    """
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    pdfs: list[Path] = []
    failures: list[str] = []
    used_names: set[str] = set()

    for raw in inputs:
        path = Path(raw)
        try:
            pdfs.extend(build_reports_for_file(path, out_dir, used_names))
        except ParseError as exc:
            failures.append(f"{path.name}: {exc}")
            log.error("%s", exc)
        except Exception as exc:                       # one bad file must not
            failures.append(f"{path.name}: {exc}")     # sink the whole batch
            log.exception("Failed to process %s", path.name)

    if not pdfs:
        raise RuntimeError(
            "No reports were generated.\n" + "\n".join(failures)
            if failures else "No reports were generated."
        )

    if len(pdfs) == 1 and not failures:
        return pdfs[0]

    zip_path = out_dir / f"{_safe_name(Path(str(inputs[0])).stem)}_reports.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as archive:
        for pdf in pdfs:
            archive.write(pdf, arcname=pdf.name)
        if failures:
            # A batch that partly failed says so inside the download, rather
            # than quietly returning fewer files than were uploaded.
            archive.writestr(
                "NOT_PROCESSED.txt",
                "The following inputs could not be processed:\n\n" + "\n".join(failures) + "\n",
            )
    log.info("Wrote %s (%s report(s))", zip_path.name, len(pdfs))
    return zip_path


class ReportStream(io.BytesIO):
    """
    The result of document_generate().

    The route names the return value `stream`, but it is not certain whether the
    previous version handed back bytes or a path, so this is both: it reads like
    a file object, it satisfies os.PathLike, and str() gives the path on disk.
    Flask's send_file accepts it either way.
    """

    def __init__(self, path: Path):
        super().__init__(path.read_bytes())
        self.path = path
        self.name = path.name
        self.mimetype = (
            "application/zip" if path.suffix == ".zip" else "application/pdf"
        )

    def __fspath__(self) -> str:
        return str(self.path)

    def __str__(self) -> str:
        return str(self.path)

    def __repr__(self) -> str:
        return f"<ReportStream {self.path.name} {len(self.getbuffer()):,} bytes>"


def _resolve_export_path(file_name: str | Path, project: str | None) -> Path:
    """
    Turn a selected file name into a path on disk.

    The route passes a project name and a file name rather than a path, and the
    convention that joins them is not recorded anywhere, so each plausible
    arrangement is tried in turn and the one that exists wins. The failure
    message lists everything that was tried, which is usually enough to identify
    the real convention on the first run.
    """
    candidate = Path(file_name)
    if candidate.is_file():
        return candidate

    tried: list[str] = [str(candidate)]
    roots = [Path(r) for r in PROJECT_ROOTS]

    layouts: list[Path] = []
    for root in roots:
        if project:
            layouts.append(root / project / candidate)
        layouts.append(root / candidate)
    if project:
        layouts.append(Path(project) / candidate)

    for layout in layouts:
        tried.append(str(layout))
        if layout.is_file():
            log.info("Resolved %s to %s", file_name, layout)
            return layout

    # Last resort: look for the filename anywhere under the roots.
    for root in roots:
        if not root.is_dir():
            continue
        for depth in range(1, PROJECT_FILE_SEARCH_DEPTH + 1):
            pattern = "/".join(["*"] * depth) + "/" + candidate.name
            for hit in root.glob(pattern):
                if hit.is_file():
                    log.info("Found %s by searching %s: %s", candidate.name, root, hit)
                    return hit

    raise FileNotFoundError(
        f"Could not find the export '{file_name}'"
        + (f" for project '{project}'" if project else "")
        + ".\nLooked in:\n  "
        + "\n  ".join(tried)
        + "\nAdd the correct directory to PROJECT_ROOTS in generate_doc.py."
    )


def document_generate(*args: Any, **kwargs: Any) -> ReportStream:
    """
    Entry point for the Flask route, kept at its original name.

        stream = document_generate(project_name_selection, file_name_selection)

    The file selection may be one name or several; several produce a ZIP. Names
    are resolved against PROJECT_ROOTS - see _resolve_export_path.

    Also accepts the other shapes this function has been called with:

        document_generate("export.xml")
        document_generate(["a.xml", "b.xml"], out_dir="reports")
        document_generate(project="P1", file_name_selection=["a.xml"])
    """
    project: Any = None
    selection: Any = None
    out_dir: Any = kwargs.get("out_dir") or kwargs.get("output_dir") or kwargs.get("dest")

    positional = list(args)
    if len(positional) >= 2:
        project, selection = positional[0], positional[1]
        if len(positional) >= 3 and out_dir is None:
            out_dir = positional[2]
        # A second string argument is ambiguous: it may be the file selection or
        # an output directory. An existing directory is treated as the latter.
        if isinstance(selection, (str, Path)) and Path(selection).is_dir():
            log.info("Second argument %r is a directory; using it as the output "
                     "directory.", str(selection))
            out_dir, selection = selection, project
            project = None
    elif positional:
        selection = positional[0]

    for key in ("project_name_selection", "project_name", "project", "project_id"):
        if project is None and key in kwargs:
            project = kwargs[key]
    for key in ("file_name_selection", "file_name", "file_names", "file_path",
                "file_paths", "files", "inputs", "xml_path", "path"):
        if selection is None and key in kwargs:
            selection = kwargs[key]

    if selection is None:
        raise TypeError(
            "document_generate() needs the export file(s) to analyse. Pass them "
            "positionally after the project name, or as file_name_selection=..."
        )

    if isinstance(selection, (str, Path)):
        selection = [selection]
    else:
        selection = list(selection)
    if not selection:
        raise ValueError("document_generate() was given an empty file selection.")

    project = str(project) if project else None
    inputs = [_resolve_export_path(name, project) for name in selection]
    log.info("Generating from %s export(s)%s.", len(inputs),
             f" in project {project}" if project else "")

    return ReportStream(generate(inputs, out_dir or DEFAULT_OUTPUT_DIR))


def main(argv: Sequence[str] | None = None) -> int:
    import argparse

    parser = argparse.ArgumentParser(
        description="Generate a workflow analysis report from DataStage XML exports."
    )
    parser.add_argument("inputs", nargs="+", help="One or more DataStage XML export files.")
    parser.add_argument("-o", "--out-dir", default=DEFAULT_OUTPUT_DIR, help="Output directory.")
    parser.add_argument("-t", "--template", default=None, help="Path to temp.docx.")
    parser.add_argument("--no-llm", action="store_true",
                        help="Skip narrative generation; produce tables only.")
    parser.add_argument("--keep-docx", action="store_true",
                        help="Keep the intermediate Word document.")
    parser.add_argument("-v", "--verbose", action="store_true")
    args = parser.parse_args(argv)

    logging.basicConfig(
        level=logging.DEBUG if args.verbose else logging.INFO,
        format="%(levelname)s  %(message)s",
    )

    global TEMPLATE_DOCX_PATH, LLM_ENABLED, KEEP_INTERMEDIATE_DOCX
    if args.template:
        TEMPLATE_DOCX_PATH = args.template
    if args.no_llm:
        LLM_ENABLED = False
    if args.keep_docx:
        KEEP_INTERMEDIATE_DOCX = True

    if not Path(TEMPLATE_DOCX_PATH).is_file():
        log.error("Template not found: %s", TEMPLATE_DOCX_PATH)
        return 1

    try:
        result = generate(args.inputs, args.out_dir)
    except RuntimeError as exc:
        log.error("%s", exc)
        return 1

    print(result)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
